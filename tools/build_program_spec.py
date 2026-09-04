from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"E:\26翌光游戏开发\策划文档\植物精灵 GGJ48H可玩Demo程序制作需求与实现规格 V1.0.docx")

FONT_CN = "Microsoft YaHei"
FONT_EN = "Aptos"
FONT_CODE = "Consolas"
BLACK = "000000"
CHARCOAL = "33413A"
PALE = "F2F6F3"
LIGHT_GRAY = "D9D9D9"
MID_GRAY = "666666"
RED = "A33A32"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=95, start=110, bottom=95, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), LIGHT_GRAY)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_run_font(run, size=None, bold=None, color=BLACK, italic=None, code=False):
    font = FONT_CODE if code else FONT_EN
    run.font.name = font
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), FONT_CN if not code else FONT_CODE)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=10.6, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=10.6)
    else:
        set_run_font(p.add_run(text), size=10.6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.16
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(item), size=10.3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12


def add_numbered(doc, items):
    for index, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.26)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.keep_together = True
        set_run_font(p.add_run(f"{index}.  {item}"), size=10.3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12


def add_code(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = 1.0
        set_run_font(p.add_run(line), size=8.9, color="2E3833", code=True)


def add_table(doc, headers, rows, widths=None, font_size=8.8, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    set_row_cant_split(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, CHARCOAL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        p.paragraph_format.keep_with_next = True
        set_run_font(p.add_run(str(value)), size=font_size, bold=True, color="FFFFFF")
    for row_index, values in enumerate(rows):
        row = table.add_row()
        set_row_cant_split(row)
        for col_index, value in enumerate(values):
            cell = row.cells[col_index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index % 2 == 1:
                set_cell_shading(cell, PALE)
            p = cell.paragraphs[0]
            p.alignment = alignments[col_index] if alignments and col_index < len(alignments) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            set_run_font(p.add_run(str(value)), size=font_size)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1, page_break=False):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    if page_break:
        p.paragraph_format.page_break_before = True
    return p


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, end])
    set_run_font(run, size=8.5, color=MID_GRAY)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.62)
section.bottom_margin = Inches(0.62)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT_EN
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
normal.font.size = Pt(10.6)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.16

title_style = styles["Title"]
title_style.font.name = FONT_EN
title_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
title_style.font.size = Pt(26)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor.from_string(BLACK)
title_ppr = title_style._element.get_or_add_pPr()
title_border = title_ppr.find(qn("w:pBdr"))
if title_border is not None:
    title_ppr.remove(title_border)

for style_name, size, before, after in (
    ("Heading 1", 17, 14, 6),
    ("Heading 2", 13, 10, 4),
    ("Heading 3", 11.2, 8, 3),
):
    style = styles[style_name]
    style.font.name = FONT_EN
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(BLACK)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for name in ("List Bullet", "List Number"):
    style = styles[name]
    style.font.name = FONT_EN
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    style.font.size = Pt(10.3)

footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(footer_p.add_run("植物精灵 GGJ48H 程序实现规格  V1.0   |   "), size=8.5, color=MID_GRAY)
add_page_field(footer_p)

# Cover
spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(35)
title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(title.add_run("植物精灵 GGJ48H 可玩 Demo 程序制作需求与实现规格"), size=26, bold=True)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(22)
set_run_font(subtitle.add_run("Unity 2022 3 62f3 程序执行版"), size=14, bold=True)

intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
intro.paragraph_format.left_indent = Inches(0.48)
intro.paragraph_format.right_indent = Inches(0.48)
intro.paragraph_format.space_after = Pt(18)
intro.paragraph_format.line_spacing = 1.24
set_run_font(
    intro.add_run(
        "程序交付目标是一条可以从主页面开始、完成第一关三段战斗、随时嫁接三种器官、清场开启传送门，并能正确处理死亡与通关的 5 至 8 分钟流程。"
        "本文件按当前 Unity 工程的实际状态制定，所有条目都包含实现边界和验收结果。"
    ),
    size=11.3,
)

add_table(
    doc,
    ["项目", "执行决定"],
    [
        ["首版核心", "移动与战斗  三种敌人  三种随时嫁接  三个战斗区  清场传送门"],
        ["正式场景", "MainMenu 和 Level01；死亡与通关使用 Level01 内的 UI 页面"],
        ["必须停用", "旧运行时灰盒生成  OnGUI 正式界面  Boss  融合  污染  进化  随机关卡"],
        ["交付平台", "Windows x86 64；1920 x 1080 为基准；键鼠必须完整可玩"],
        ["完成标准", "脱离 Unity 编辑器启动后，连续重玩 3 次均无软锁 无重复结算 无残留状态"],
    ],
    widths=[1.4, 5.7],
    font_size=9.4,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(12)
set_run_font(meta.add_run("版本 1.0  更新日期 2026 年 9 月 4 日"), size=9.5, color=MID_GRAY)

doc.add_page_break()

add_heading(doc, "1 当前工程诊断与重做结论", 1)
add_body(
    doc,
    "当前项目可以证明部分规则，但还不是可提交的游戏工程。核心问题不是代码数量少，而是正式内容仍由运行时代码临时生成，旧版与新版逻辑同时编译，且关卡、界面、嫁接效果和流程状态没有稳定的资产与职责边界。继续在现有大脚本内增加内容，会让美术接入、流程修复和最终构建变得更慢。",
)

add_heading(doc, "1.1 已确认的当前状态", 2)
add_table(
    doc,
    ["检查项", "当前事实", "造成的问题", "处理决定"],
    [
        ["Unity 版本", "2022.3.62f3", "版本本身可用", "全队锁定该版本，不在比赛中升级"],
        ["Build Settings", "场景列表为空", "无法形成明确启动入口和发布流程", "创建 MainMenu 与 Level01 并加入构建顺序"],
        ["启动方式", "RuntimeInitializeOnLoad 自动创建全部对象", "任何空场景都可能生成游戏，难以定位重复实例", "改为场景中的 GameBootstrap 负责初始化"],
        ["关卡", "BuildGreybox 在代码中创建平台", "美术无法直接编辑 Tilemap 和碰撞", "Level01 使用正式场景层级与 Tilemap"],
        ["UI", "GgjUi 和 VerticalSliceHud 使用 OnGUI 固定坐标", "分辨率变化、锚点和美术替换均不可控", "改用 Canvas  CanvasScaler  Prefab 和 Presenter"],
        ["数据", "VerticalSliceCatalog 静态写死", "调数值必须改代码，仍包含旧范围内容", "首版只建立 3 敌人和 3 嫁接的 ScriptableObject"],
        ["旧系统", "Boss 融合 污染 进化 随机房间仍在编译", "与当前首版边界冲突，增加测试分支", "移出 Assets 或从 GGJ 程序集中排除"],
        ["角色死亡", "当前直接 Respawn 并回满生命", "没有死亡结算，也无法验证重开状态", "死亡进入 Dead 状态，显示死亡页后由玩家选择"],
        ["嫁接效果", "铁甲根无挡弹，毒雾无减速，藤蔓判定近似", "展示内容与策划承诺不一致", "按本规格实现并加入单项验收场景"],
        ["流程计数", "全局 activeEnemies 可被后续区域覆盖", "越过触发线可能提前开下一战斗区或错误清场", "每个 Encounter 独立维护敌人集合并锁门"],
    ],
    widths=[1.0, 1.8, 2.2, 2.1],
    font_size=8.15,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "1.2 重做边界", 2)
add_bullets(
    doc,
    [
        "保留可验证的基础数值和命名，但不继续维护 PlantSpiritDemo 与 VerticalSlice 两套并行流程。",
        "移除 VerticalSliceRuntime 上的 RuntimeInitializeOnLoadMethod。旧代码先在版本控制中打标签，再移动到项目根目录 Archive，不继续参与 Unity 编译。",
        "首版不建设通用技能树、通用 Buff 框架、依赖注入容器、Addressables、网络或复杂存档。每项抽象必须直接服务本 Demo。",
        "正式玩法只从 MainMenu 进入 Level01。程序不再依赖打开任意空场景即可运行的调试方式。",
        "所有功能先用正式 Prefab 和占位 Sprite 跑通，再由美术在不改脚本引用的前提下替换 Sprite、Animator 和 VFX。",
    ],
)

delayed_sections_start = len(doc.element.body) - 1
add_heading(doc, "5 输入系统与操作合同", 1)
add_body(doc, "输入必须集中在 InputReader。玩家、UI、传送门和嫁接界面只能订阅语义事件，不能在各自 Update 中散落 Input.GetKey。这样才能在暂停、嫁接、死亡和换场景时准确关闭输入。")

add_heading(doc, "5.1 输入方案", 2)
add_table(
    doc,
    ["项目", "决定", "完成标准"],
    [
        ["输入包", "安装与 Unity 2022.3 LTS 兼容的 Input System，并启用 Both 作为过渡", "工程重启后无输入后端警告"],
        ["Action Asset", "PlantSpirit.inputactions", "只有 Gameplay 和 UI 两个 Action Map"],
        ["键鼠", "P0 必须完整", "不用手柄即可通关 重开 返回主页面"],
        ["手柄", "P1", "不影响键鼠提交；若实现则使用标准 Gamepad 路径"],
        ["订阅生命周期", "OnEnable 订阅  OnDisable 退订", "场景重载 3 次后一次按键只触发一次"],
    ],
    widths=[1.25, 3.35, 2.5],
    font_size=8.7,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "5.2 Action 映射", 2)
add_table(
    doc,
    ["Action", "类型", "键鼠", "Gameplay 状态结果"],
    [
        ["Move", "Value Vector2", "A D 和左右方向键", "只读取 X；更新角色朝向"],
        ["Jump", "Button", "Space", "按下写入 0.12 秒跳跃缓冲"],
        ["Dash", "Button", "Left Shift", "满足冷却和状态条件时开始冲刺"],
        ["Attack", "Button", "鼠标左键或 J", "写入 0.12 秒攻击缓冲"],
        ["Skill", "Button", "鼠标右键或 K", "技能冷却结束时施放当前花部技能"],
        ["Interact", "Button", "E", "仅作用于当前最近的可交互对象"],
        ["Graft", "Button", "Tab 或 G", "Playing 打开嫁接；Grafting 关闭嫁接"],
        ["Pause", "Button", "Esc", "Playing 打开暂停；Paused 继续；Grafting 关闭嫁接"],
    ],
    widths=[1.0, 1.05, 1.7, 3.35],
    font_size=8.45,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "5.3 输入优先级与保护", 2)
add_bullets(
    doc,
    [
        "打开 UI 的同一帧立即禁用 Gameplay Map，避免 Tab 或确认键继续传给角色。",
        "关闭嫁接后等待 0.20 秒 unscaled time 再启用 Attack 与 Skill；Move 可以立即恢复。",
        "跳跃和攻击使用各自输入缓冲；冲刺、技能和交互不缓存，避免冷却结束后自动触发。",
        "角色处于 Hurt 或 Dead 时忽略 Attack Skill Dash。处于 Dash 时忽略 Jump 和 Attack。",
        "UI 的 Submit 和 Cancel 必须与 Gameplay 的 Attack 和 Pause 隔离，不共用同一个回调对象。",
    ],
)

add_heading(doc, "6 玩家移动实现", 1)
add_body(doc, "角色使用 Rigidbody2D 动态刚体和 Collider2D。Update 收集输入并推进状态计时，FixedUpdate 只处理速度与物理。禁止直接修改 transform.position 完成普通移动或冲刺。")

add_heading(doc, "6.1 Player Prefab 组件", 2)
add_table(
    doc,
    ["组件", "关键设置", "职责"],
    [
        ["Rigidbody2D", "Dynamic  Freeze Rotation Z  Interpolate  Continuous", "地面碰撞 跳跃 冲刺和击退"],
        ["CapsuleCollider2D", "贴合躯干，脚底略高于 Sprite 底部", "实体碰撞，不作为受击范围"],
        ["PlayerMotor2D", "引用 Rigidbody GroundCheck PlayerConfig", "移动 跳跃 冲刺 地面判定"],
        ["PlayerCombat", "引用攻击挂点 Hitbox 技能发射点", "执行当前茎部普攻和花部技能"],
        ["Health", "maxHealth 100", "生命 受伤 无敌 死亡事件"],
        ["Hurtbox2D", "PlayerHurtbox Layer", "接收 DamageInfo"],
        ["GraftLoadout", "Root Stem Flower", "保存当前三个槽位并提供查询"],
        ["PlayerAnimationBridge", "Animator SpriteRenderer 挂点", "状态参数 外观刷新 受击闪烁"],
    ],
    widths=[1.55, 2.7, 2.85],
    font_size=8.55,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "6.2 移动参数", 2)
add_table(
    doc,
    ["字段", "推荐值", "实现规则", "验收"],
    [
        ["groundMaxSpeed", "5.2", "MoveTowards 逼近目标速度", "平地 2 秒位移约 10.4 单位"],
        ["groundAcceleration", "42", "只在有输入时使用", "起步轻快，不瞬间满速"],
        ["groundDeceleration", "55", "松开输入时使用", "松键后约 0.10 秒停稳"],
        ["airAcceleration", "28", "空中可修正但弱于地面", "跳跃中可以回拉但不能急停"],
        ["jumpVelocity", "12.5", "满足 Grounded 或 Coyote 时写入 Y 速度", "可稳定越过设计高差"],
        ["coyoteTime", "0.12 秒", "离地后保留跳跃资格", "刚离平台仍可起跳"],
        ["jumpBuffer", "0.12 秒", "落地前按键自动在落地帧起跳", "不重复消耗"],
        ["maxFallSpeed", "18", "限制负向 Y 速度", "下落保持可控"],
        ["dashDistance", "3.6", "由速度和 0.22 秒时长共同得到", "不能穿过 Ground 碰撞"],
        ["dashDuration", "0.22 秒", "状态计时结束即恢复普通移动", "帧率变化不改变距离超过 5%"],
        ["dashCooldown", "1.10 秒", "从冲刺开始计时", "冷却未结束不能重入"],
        ["dashInvulnerable", "0.18 秒", "只覆盖前段", "后 0.04 秒可受击"],
    ],
    widths=[1.35, 1.15, 2.8, 1.8],
    font_size=8.1,
    alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "6.3 玩家状态优先级", 2)
add_table(
    doc,
    ["优先级", "状态", "可转入条件", "锁定内容", "退出条件"],
    [
        ["1", "Dead", "Health 首次降至 0", "全部玩法输入和伤害", "重开或返回主页面"],
        ["2", "Hurt", "收到有效伤害且未无敌", "攻击 技能 冲刺", "0.28 秒后回到 Air 或 Ground"],
        ["3", "Dash", "冷却结束且非 Hurt Attack", "普通移动 跳跃 攻击", "0.22 秒或撞墙"],
        ["4", "Attack", "攻击缓冲存在且非 Dash Hurt", "冲刺和技能；可保留轻微水平惯性", "攻击后摇结束"],
        ["5", "Air", "GroundCheck 失去地面", "无", "落地或更高优先级"],
        ["6", "Ground", "GroundCheck 有效", "无", "离地或更高优先级"],
    ],
    widths=[0.7, 0.85, 2.0, 2.0, 1.55],
    font_size=8.25,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "6.4 地面和边界规则", 2)
add_bullets(
    doc,
    [
        "GroundCheck 使用脚底 BoxCast 或 OverlapBox，只检测 Ground Layer；不能把敌人或掉落物当作地面。",
        "落地只在法线方向合理且垂直速度小于阈值时成立，避免贴墙刷新土狼时间。",
        "冲刺前用 Rigidbody2D.Cast 检查路径，按最近墙体距离截短，不允许穿墙或卡入 Tilemap。",
        "玩家掉到 KillPlane 后扣除 20 生命并回到最近战斗区安全点；若扣血致死则正常进入 Dead。",
        "镜头边界不代替物理边界。Level01 左右端必须有不可见 Ground Collider。",
    ],
)

add_heading(doc, "7 生命 伤害与命中系统", 1)
add_body(doc, "敌我双方共用 DamageInfo、Hitbox2D、Hurtbox2D 和 Health。攻击脚本只产生伤害请求，Health 决定是否接受并发出反馈事件。所有伤害必须带攻击实例编号，用于防止一个攻击窗口重复命中同一目标。")

add_heading(doc, "7.1 核心数据合同", 2)
add_code(
    doc,
    [
        "DamageInfo",
        "  float amount",
        "  Vector2 hitPoint",
        "  Vector2 knockback",
        "  GameObject source",
        "  int attackInstanceId",
        "  DamageType type  Physical or Poison",
        "  bool canBreakProjectile",
        "Health.TryDamage DamageInfo returns bool",
        "Hitbox2D.Open attackInstanceId duration maxTargets",
        "Hitbox2D.Close",
    ],
)

add_heading(doc, "7.2 伤害处理顺序", 2)
add_numbered(
    doc,
    [
        "Hurtbox 收到 Hitbox 或 Projectile 的 DamageInfo，先检查双方是否仍在有效 Gameplay 状态。",
        "检查 attackInstanceId 是否已经命中过该 Hurtbox；重复则返回 false。",
        "检查目标死亡、无敌、阵营和护盾。铁甲根冲刺护盾可先消费普通敌方投射物。",
        "计算最终伤害。铁甲根把物理与毒伤害乘以 0.75，结果保留一位小数；UI 向上取整显示生命。",
        "扣除生命并触发 Damaged 事件，事件包含最终伤害、命中点和击退。",
        "生命降至 0 时立即置为 Dead，清空碰撞伤害和输入；死亡事件只允许发出一次。",
    ],
)

add_heading(doc, "7.3 命中反馈时序", 2)
add_table(
    doc,
    ["事件", "程序触发", "表现监听", "硬规则"],
    [
        ["普通命中", "Health.Damaged", "受击闪色  粒子  音效  0.04 秒 HitStop", "HitStop 使用独立服务，结束后恢复原 timeScale"],
        ["重命中", "DamageInfo type 或 AttackDefinition 标记", "更强震屏与 0.07 秒 HitStop", "同一帧多目标只触发一次全局停顿"],
        ["玩家受击", "PlayerHealth.Damaged", "闪红  击退  受击声  血条动画", "0.75 秒受击无敌"],
        ["挡弹", "IronRootShield.Blocked", "木甲碎裂  独立声音", "投射物立即失效，不再造成伤害"],
        ["死亡", "Health.Died", "死亡动画  清除伤害判定", "区域计数立即减少，尸体稍后销毁"],
    ],
    widths=[1.1, 1.65, 2.6, 1.75],
    font_size=8.35,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "7.4 性能和可靠性要求", 2)
add_bullets(
    doc,
    [
        "攻击期间使用预分配 Collider2D 数组或 ContactFilter2D，不在每帧创建 List。",
        "同一攻击的已命中目标集合在攻击结束后清空，不能跨攻击保留。",
        "投射物数量少时允许 Instantiate 和 Destroy，但必须有 3 秒超时、关卡边界销毁和场景清理。若数量超过 20 再加入小型对象池。",
        "敌人死亡逻辑先从 Encounter 注销，再播放动画。Destroy 延迟不能阻塞清场。",
        "暂停和嫁接期间不运行伤害计时；毒伤、敌人冷却和投射物全部使用 scaled deltaTime。",
    ],
)

add_heading(doc, "8 玩家攻击与三种嫁接", 1, page_break=True)
add_body(doc, "首版不实现通用技能编辑器。PlayerCombat 固定支持四种执行器：基础近战、藤蔓线性近战、种子投射物和毒雾区域。GraftLoadout 只决定使用哪一种执行器及其配置，减少比赛期间动态组件和残留状态问题。")

add_heading(doc, "8.1 默认能力", 2)
add_table(
    doc,
    ["能力", "数值", "判定", "执行时序"],
    [
        ["嫩芽抽打", "10 伤害  最短间隔 0.45 秒", "面向方向 1.2 x 1.0 矩形  最多 1 目标", "前摇 0.12  激活 0.08  后摇 0.18"],
        ["种子射击", "8 伤害  冷却 3 秒", "速度 8.5  半径 0.18  寿命 1.2 秒", "发射前摇 0.10  生成后进入冷却"],
        ["基础冲刺", "无伤害  冷却 1.10 秒", "路径碰墙截短  前 0.18 秒无敌", "0.22 秒结束，保留少量水平速度"],
    ],
    widths=[1.15, 1.7, 2.4, 1.85],
    font_size=8.5,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "8.2 GraftDefinition 数据字段", 2)
add_table(
    doc,
    ["字段", "类型", "用途", "首版约束"],
    [
        ["id", "string", "稳定唯一 ID", "vine_tendril  toxic_cap  iron_root"],
        ["displayName", "string", "UI 名称", "藤蔓触须  毒菌伞  铁甲根"],
        ["slot", "GraftSlot", "Root Stem Flower", "每个部件只属于一个槽位"],
        ["icon", "Sprite", "掉落提示与嫁接界面", "缺失时数据校验报错"],
        ["equippedVisual", "GameObject", "角色挂点外观 Prefab", "不能自带 gameplay 脚本"],
        ["abilityType", "AbilityType", "选择固定执行器", "None MeleeBox VineLine Projectile PoisonCloud IronRoot"],
        ["attackData", "AttackDefinition", "伤害范围节奏和 Prefab", "Root 可为空，Stem Flower 必填"],
        ["damageReduction", "float", "根部减伤", "铁甲根 0.25，其余 0"],
        ["blocksProjectileOnDash", "bool", "冲刺护盾", "只有铁甲根为 true"],
        ["description", "string", "UI 说明", "必须直接说明行为变化"],
    ],
    widths=[1.55, 1.25, 2.55, 1.75],
    font_size=8.2,
    alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "8.3 藤蔓触须实现", 2)
add_table(
    doc,
    ["项目", "要求"],
    [
        ["装备槽位", "Stem；替换嫩芽抽打，不影响花部技能和根部被动"],
        ["攻击数值", "8 伤害；最短间隔 0.70 秒；前摇 0.14 秒；有效 0.10 秒；后摇 0.22 秒"],
        ["判定", "从 StemAttackPoint 向面向方向 BoxCast 2.8 单位，宽 0.65；按距离排序，最多命中 2 个不同敌人"],
        ["穿透定义", "最多命中前方两个敌人。不是射线无限穿透，也不能命中角色身后的敌人"],
        ["反馈", "长藤鞭轨迹  叶片碎屑  挥鞭声；第二个目标也有受击反馈，但全局 HitStop 只触发一次"],
        ["验收摆位", "两个敌人放在 1.6 和 2.5 单位可同时命中；第三个 2.7 单位敌人不扣血；身后敌人不扣血"],
    ],
    widths=[1.35, 5.75],
    font_size=8.75,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "8.4 毒菌伞实现", 2)
add_table(
    doc,
    ["项目", "要求"],
    [
        ["装备槽位", "Flower；替换种子射击，不改变普通攻击"],
        ["技能数值", "冷却 5 秒；毒雾持续 3 秒；每 0.5 秒结算 3 伤害，总计最多 18 伤害；移速降低 30%"],
        ["生成范围", "角色前方 0.8 单位生成 PoisonZone，建议 3.0 x 1.8 矩形或等效扇形 Collider"],
        ["结算规则", "进入范围立即挂 Poison 状态；重复进入刷新剩余持续时间，不叠加多个 DPS；同一毒雾不可重复创建状态实例"],
        ["减速规则", "只修改敌人的移动速度倍率。藤蔓追踪和蘑菇短移生效；甲虫蓄力可减速，冲刺开始后免疫减速"],
        ["离开与销毁", "敌人离开雾区后已挂的毒状态继续到 3 秒结束；PoisonZone 3 秒后销毁，场景退出强制清理"],
        ["验收", "技能不再生成种子；单个敌人 3 秒共受 6 次毒伤；重新施放只刷新，不把 DPS 翻倍"],
    ],
    widths=[1.35, 5.75],
    font_size=8.7,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "8.5 铁甲根实现", 2)
add_table(
    doc,
    ["项目", "要求"],
    [
        ["装备槽位", "Root；不替换普攻和花部技能"],
        ["减伤", "最终伤害乘以 0.75。甲虫 12 点撞击应变为 9 点；不使用整数截断"],
        ["冲刺护盾", "冲刺开始至 0.18 秒结束前启用 ShieldHurtbox，只检测 EnemyProjectile Layer"],
        ["挡弹", "普通蘑菇毒球接触护盾后立即播放挡弹反馈并销毁，不触发玩家 Hurtbox；甲虫本体冲撞不能被护盾销毁"],
        ["外观", "RootVisualPoint 替换为铁灰根系外观；冲刺期间额外显示护盾轮廓，结束立即隐藏"],
        ["卸下", "减伤恢复 0，ShieldHurtbox 永久关闭并清理残留事件订阅"],
        ["验收", "装备与卸下分别受到三次 12 伤害，扣血结果稳定为 9 和 12；冲刺前段挡弹，后段仍会受击"],
    ],
    widths=[1.35, 5.75],
    font_size=8.7,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "8.6 嫁接事务顺序", 2)
add_numbered(
    doc,
    [
        "GraftPresenter 只能提交库存中存在且槽位匹配的 GraftDefinition；无效请求不关闭界面并显示错误。",
        "GraftApplier 保存旧槽位引用，停止旧能力仍在运行的对象，例如未发射的攻击窗口或旧护盾。",
        "GraftLoadout 写入新部件，并从 PlayerConfig 基础值重新计算最终属性，不能在旧数值上连续加减。",
        "PlayerCombat 根据 abilityType 刷新执行器和 AttackDefinition；PlayerAnimationBridge 替换对应挂点外观。",
        "HudPresenter 刷新槽位图标与技能冷却显示，然后发出 GraftApplied 反馈事件。",
        "任一步骤抛出异常时恢复旧槽位，记录部件 ID 和失败环节；首版界面保持打开，避免角色进入半装备状态。",
    ],
)

add_heading(doc, "8.7 嫁接界面行为", 2)
add_table(
    doc,
    ["步骤", "界面结果", "程序要求"],
    [
        ["打开", "游戏冻结，默认选中最近拾取部件", "State 切换 Grafting；刷新快照；不每帧遍历场景"],
        ["选槽", "左侧根 茎 花高亮", "右侧仅显示该槽位库存；基础能力也可选择"],
        ["选部件", "显示当前与新效果对比", "预览不修改实际角色和冷却"],
        ["确认", "播放嫁接反馈并装备", "一次点击只提交一次；成功后关闭界面"],
        ["取消", "返回战斗，不改变装备", "恢复 Playing 并启用 0.20 秒输入保护"],
        ["空库存", "显示尚未获得该部件", "不创建空按钮或空引用"],
    ],
    widths=[1.0, 2.5, 3.6],
    font_size=8.65,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

# This block is authored here for maintainability, then placed after section 4.
delayed_sections = list(doc.element.body)[delayed_sections_start:-1]
for element in delayed_sections:
    doc.element.body.remove(element)

add_heading(doc, "2 程序交付范围与完成定义", 1)
add_body(doc, "程序必须交付的是完整可玩闭环，不是若干独立演示。以下 P0 项目缺少任何一项都不算首版完成。")

add_heading(doc, "2.1 P0 交付物", 2)
add_table(
    doc,
    ["编号", "交付物", "玩家能看到的结果", "程序完成标准"],
    [
        ["P0 01", "主页面", "开始游戏  操作  音量  退出", "按钮仅触发一次；返回主页面后可再次开局"],
        ["P0 02", "第一关", "教学起点  三个战斗区  出口", "区域顺序固定；未清场不能穿过封锁门"],
        ["P0 03", "角色控制", "移动 跳跃 冲刺 朝向", "手感参数可配置；暂停和死亡时不接收 Gameplay 输入"],
        ["P0 04", "角色战斗", "普攻  默认种子技能  受击  死亡", "伤害不重复；冷却正确；死亡只结算一次"],
        ["P0 05", "三种敌人", "藤蔓近战  蘑菇远程  甲虫冲刺", "每种敌人有前摇 攻击 后摇 受击 死亡和防卡死规则"],
        ["P0 06", "掉落与拾取", "三个关键器官必定获得", "指定敌人首杀必掉；拾取不会重复加入；重开清空"],
        ["P0 07", "随时嫁接", "战斗中打开界面，确认后能力与外观立即变化", "暂停完全；关闭有输入保护；旧效果无残留"],
        ["P0 08", "三种嫁接", "藤蔓改普攻  毒菌改技能  铁甲改防御与冲刺", "数值、判定、VFX 与 UI 同步变化"],
        ["P0 09", "传送门", "最后一个区域清场后生长并开启", "未清场不可用；进入只触发一次并显示结果页"],
        ["P0 10", "发布构建", "Windows 包可直接运行", "第二台无 Unity 电脑完整通关 3 次"],
    ],
    widths=[0.7, 1.3, 2.5, 2.6],
    font_size=8.25,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "2.2 明确不进入首版的内容", 2)
add_table(
    doc,
    ["内容", "首版处理", "以后接入位置"],
    [
        ["融合", "不显示入口，不保留自动触发逻辑", "后续在关卡中的融合台打开独立界面"],
        ["Boss", "不创建 Prefab，不进入流程", "第二阶段在传送门之后增加 Boss 关"],
        ["随机房间", "第一关为手工固定场景", "EncounterDefinition 可复用后再做房间池"],
        ["污染与进化", "删除运行时字段和 UI", "核心战斗验证后再建立独立系统"],
        ["永久成长和存档", "只保存音量；不保存单局进度", "完成首章结构后再设计版本迁移"],
        ["手柄完整支持", "键鼠为 P0，手柄为 P1", "InputAction 已预留绑定，但不影响提交"],
    ],
    widths=[1.55, 2.8, 2.75],
    font_size=8.8,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "2.3 首版完成定义", 2)
add_numbered(
    doc,
    [
        "从 Windows 可执行文件进入主页面，点击开始后稳定载入 Level01。",
        "玩家不阅读外部说明也能沿场景向右推进，完成三次战斗并获得三个器官。",
        "三个器官都能在战斗中打开嫁接界面装备，且各自改变真实战斗结果。",
        "最后一只敌人死亡后，所有敌方攻击停止，传送门开启；玩家进入后显示用时和已装备器官。",
        "死亡后可重开，通关后可再次挑战或返回主页面；三次连续循环无静态实例、订阅、库存或 Time.timeScale 残留。",
        "Development Build 和普通 Build 均无 Error 日志；允许记录明确的非阻断 Warning，但提交前必须归档原因。",
    ],
)

add_heading(doc, "3 游戏运行状态与流程合同", 1)
add_body(doc, "所有页面、输入、暂停和结算必须由一个 GameStateController 管理。各脚本不能自行改 Time.timeScale 或加载场景，否则会出现嫁接关闭后仍暂停、死亡时还能攻击、重复进入传送门等问题。")

add_heading(doc, "3.1 全局状态", 2)
add_table(
    doc,
    ["状态", "允许输入", "时间", "进入动作", "允许离开到"],
    [
        ["MainMenu", "UI", "1", "显示主页面，Gameplay Map 关闭", "Loading"],
        ["Loading", "无", "1", "屏蔽按钮，异步加载 Level01", "Playing"],
        ["Playing", "Gameplay", "1", "恢复玩家、AI、投射物", "Paused  Grafting  Dead  Result"],
        ["Paused", "UI", "0", "打开暂停页，保存此前为 Playing", "Playing  MainMenu"],
        ["Grafting", "UI", "0", "打开嫁接页，刷新库存并锁 Gameplay", "Playing"],
        ["Dead", "UI", "先 1 后 0", "播放死亡约 0.8 秒，再打开死亡页", "Loading  MainMenu"],
        ["Result", "UI", "0", "记录通关用时和装备，打开结果页", "Loading  MainMenu"],
    ],
    widths=[0.85, 1.0, 0.55, 2.8, 1.9],
    font_size=8.2,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "3.2 状态切换硬规则", 2)
add_bullets(
    doc,
    [
        "只有 GameStateController.SetState 可以修改全局状态和 Time.timeScale。重复请求当前状态直接返回。",
        "Paused 与 Grafting 不能互相覆盖。处于 Grafting 时按 Esc 只关闭嫁接，不能再叠加暂停页。",
        "Dead 和 Result 为终止态。进入后拒绝攻击、嫁接、暂停、受伤、区域触发和传送门交互。",
        "场景加载前先把 Time.timeScale 恢复为 1，清理静态事件，再执行 SceneManager.LoadSceneAsync。",
        "运行计时只在 Playing 累加 Time.deltaTime；暂停和嫁接时间不计入通关用时。",
        "所有延迟结算协程在对象禁用或场景退出时取消，避免旧场景协程在新局触发。",
    ],
)

add_heading(doc, "3.3 第一关事件顺序", 2)
add_numbered(
    doc,
    [
        "MainMenu 的开始按钮发出 StartRun，GameSession 清空本局数据并加载 Level01。",
        "LevelFlow 注册三个 EncounterZone、ExitPortal 和玩家出生点，然后切换到 Playing。",
        "玩家进入战斗区触发器，区域关闭左右门，生成或启用本区域敌人，并显示剩余数量。",
        "本区域每名敌人死亡时从该区域存活集合移除。集合首次变为空时区域进入 Cleared，打开门并结算必掉器官。",
        "玩家拾取器官后，GraftInventory 解锁该部件。按 Tab 或 G 进入 Grafting，确认装备后回到 Playing。",
        "第三战斗区 Cleared 后，LevelFlow 清理敌方投射物，调用 ExitPortal.BeginOpen。",
        "玩家进入传送门交互范围并按 E，Portal 进入 Entering，锁定角色输入，播放 0.6 秒吸入或淡出后进入 Result。",
    ],
)

add_heading(doc, "4 正式工程结构", 1)
add_body(doc, "首版使用少量明确脚本和 Prefab。场景对象负责摆放，ScriptableObject 负责数值，MonoBehaviour 负责运行行为。禁止再次用一个启动脚本在 Awake 中创建整关。")

add_heading(doc, "4.1 场景与构建顺序", 2)
add_table(
    doc,
    ["序号", "场景", "主要内容", "加载规则"],
    [
        ["0", "MainMenu", "GameBootstrap  EventSystem  MenuCanvas  AudioService", "应用启动入口；PersistentSystems 只创建一次"],
        ["1", "Level01", "Tilemap  Player  Camera  3 Encounter  Portal  GameCanvas", "每次重玩整场重载，避免局内对象残留"],
    ],
    widths=[0.65, 1.25, 3.4, 1.8],
    font_size=8.8,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "4.2 推荐目录", 2)
add_table(
    doc,
    ["目录", "必须存放的内容"],
    [
        ["Assets Game Scenes", "MainMenu.unity  Level01.unity"],
        ["Assets Game Scripts Core", "GameBootstrap  GameSession  GameStateController  SceneLoader"],
        ["Assets Game Scripts Input", "InputReader  PlantSpirit.inputactions"],
        ["Assets Game Scripts Player", "PlayerMotor2D  PlayerCombat  PlayerHealth  PlayerAnimationBridge"],
        ["Assets Game Scripts Combat", "DamageInfo  Hitbox2D  Hurtbox2D  Projectile2D  PoisonZone  HitStop"],
        ["Assets Game Scripts Grafting", "GraftInventory  GraftLoadout  GraftApplier  GraftDefinition"],
        ["Assets Game Scripts Enemies", "EnemyController  VineEnemy  MushroomEnemy  BeetleEnemy"],
        ["Assets Game Scripts Level", "LevelFlow  EncounterZone  EncounterDefinition  ExitPortal  KillPlane"],
        ["Assets Game Scripts UI", "MenuPresenter  HudPresenter  GraftPresenter  PausePresenter  EndPresenter"],
        ["Assets Game Data", "Player  Attack  Enemy  Graft  Encounter  Audio 配置资产"],
        ["Assets Game Prefabs", "Player  Enemies  Attacks  Drops  Portal  UI  VFX"],
        ["Assets Game Tests", "EditMode 数据校验和 PlayMode 主流程测试"],
    ],
    widths=[2.4, 4.7],
    font_size=8.6,
    alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "4.3 场景层级", 2)
add_code(
    doc,
    [
        "MainMenu",
        "  PersistentSystems  GameBootstrap  GameSession  GameStateController  AudioService",
        "  EventSystem",
        "  MenuCanvas  Background  Title  StartButton  ControlsButton  AudioPanel  QuitButton",
        "Level01",
        "  LevelRoot  LevelFlow  SpawnPoint  KillPlane",
        "  Environment  Grid  BackgroundTilemap  GroundTilemap  ForegroundTilemap",
        "  Encounters  Encounter01  Encounter02  Encounter03",
        "  ExitPortal",
        "  Player",
        "  MainCamera",
        "  GameCanvas  HUD  GraftPanel  PausePanel  DeadPanel  ResultPanel  InteractionPrompt",
    ],
)

add_heading(doc, "4.4 程序程序集与旧代码隔离", 2)
add_bullets(
    doc,
    [
        "创建 PlantSpirit.Game.asmdef，正式脚本只引用 UnityEngine UI TextMeshPro 和 Input System。",
        "把 Assets VerticalSlice 与 Assets Scripts PlantSpiritDemo.cs 移到项目根目录 Archive，或删除 RuntimeInitializeOnLoad 入口并从正式程序集排除。",
        "正式命名空间统一为 PlantSpirit.Game。禁止在同一工程中继续使用 PlantSpirit.VerticalSlice 的单例和全局列表推进新流程。",
        "场景引用全部通过 Inspector 明确绑定。Awake 只校验引用和缓存组件，不创建关卡、相机、玩家、敌人或 UI。",
    ],
)

for element in delayed_sections:
    doc.element.body.insert(len(doc.element.body) - 1, element)

add_heading(doc, "9 三种敌人实现", 1)
add_body(doc, "敌人共用 EnemyController、Health、Hurtbox2D 和状态基类，每个品种只实现自己的决策与攻击。首版不做 NavMesh、不做跨平台追路；敌人只在所属战斗区的地面段活动，遇到断崖或区域边界就停下或返回。")

add_heading(doc, "9.1 Enemy Prefab 通用组件", 2)
add_table(
    doc,
    ["组件", "职责", "硬性要求"],
    [
        ["EnemyController", "当前状态 目标 计时和区域归属", "Disable 时退订 Encounter；死亡不可重复上报"],
        ["Rigidbody2D", "移动 撞墙 击退", "Dynamic  Freeze Rotation Z；只与 Ground 碰撞"],
        ["Collider2D", "实体与墙体碰撞", "不兼任攻击判定"],
        ["Hurtbox2D", "接收玩家伤害", "EnemyHurtbox Layer；死亡立即禁用"],
        ["EnemyAttack", "生成近战 Hitbox 或投射物", "只在 Attack Active 时开启"],
        ["GroundAheadCheck", "检测前方地面和战斗区边界", "无地面时停止，不自行跳跃"],
        ["Animator", "移动 前摇 攻击 受击 死亡", "动画时长必须与数据时序一致"],
        ["AudioEmitter", "局部攻击与受击音效", "由事件调用，不在 Update 播放"],
    ],
    widths=[1.6, 2.6, 2.9],
    font_size=8.5,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "9.2 通用状态和转换", 2)
add_table(
    doc,
    ["状态", "进入条件", "行为", "离开条件"],
    [
        ["Dormant", "区域未开始", "禁用 AI 与攻击，保留显示或保持隐藏", "Encounter 激活"],
        ["Idle", "激活后无目标或短暂停顿", "站立，扫描玩家", "玩家进入感知范围"],
        ["Chase", "目标有效且不在攻击条件", "向玩家移动并检查边界", "进入攻击范围或需要返回"],
        ["Telegraph", "攻击冷却结束且满足位置条件", "锁定方向，播放前摇和预警", "前摇完成或被死亡打断"],
        ["Attack", "前摇完成", "只开启一次 Hitbox 或生成一次投射物", "有效时间结束"],
        ["Recover", "攻击结束", "不能再次攻击，可按品种缓慢移动或停顿", "后摇和冷却条件满足"],
        ["Stunned", "甲虫撞墙或特殊打断", "停止移动和攻击", "眩晕时间结束"],
        ["Hurt", "收到有效伤害且未死亡", "短闪色与击退", "0.12 至 0.20 秒后回到原决策"],
        ["Return", "超出出生区或失去地面", "回到安全点，不攻击", "进入出生点容差"],
        ["Dead", "Health 为 0", "注销计数，禁用判定，播放死亡", "0.6 至 1.0 秒后销毁"],
    ],
    widths=[1.0, 2.0, 2.4, 1.7],
    font_size=8.15,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "9.3 腐化藤蔓怪", 2)
add_table(
    doc,
    ["项目", "规格"],
    [
        ["基础数值", "生命 25；移动 1.8；攻击伤害 8；攻击距离 1.2；完整冷却 1.5 秒"],
        ["感知", "同战斗区内始终知道玩家位置；不跨区域门追踪"],
        ["追踪", "保持脚底地面检测。距离大于 1.2 时靠近；距离小于 0.8 时不继续挤压玩家"],
        ["前摇", "0.45 秒；进入时锁定面向；身体后仰，藤蔓轮廓伸长并播放短提示声"],
        ["攻击", "前方 1.35 x 0.9 矩形 Hitbox，有效 0.10 秒；一次攻击最多命中玩家一次"],
        ["后摇", "0.35 秒不能移动；之后进入冷却。玩家离开范围不会取消已经开始的攻击"],
        ["掉落", "Encounter01 的首只藤蔓死亡必掉藤蔓触须；其他藤蔓首版不随机掉落"],
        ["防卡死", "离出生点超过 5.5 单位或 Y 下降超过 1.5 时进入 Return；3 秒回不去则传回安全点"],
    ],
    widths=[1.35, 5.75],
    font_size=8.65,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "9.4 毒孢蘑菇", 2)
add_table(
    doc,
    ["项目", "规格"],
    [
        ["基础数值", "生命 18；短移速度 0.6；毒球伤害 6；射程 6.0；完整冷却 2.5 秒"],
        ["站位", "优先保持出生点附近，不追过台阶；玩家距离小于 2.2 时向后短移，最多 1.5 单位"],
        ["瞄准", "前摇开始时记录玩家胸口位置，不持续追踪；预警线或菌盖朝向表示发射方向"],
        ["前摇", "0.65 秒；菌盖收缩再张开；被普通受击不取消，被死亡取消"],
        ["投射物", "速度 4.5；寿命 3 秒；使用 EnemyProjectile Layer；碰 Ground PlayerShield PlayerHurtbox 后销毁"],
        ["后摇", "0.25 秒，随后等待剩余冷却"],
        ["掉落", "Encounter02 的指定蘑菇死亡必掉毒菌伞"],
        ["防卡死", "投射物离开关卡边界或相机扩展边界立即销毁；蘑菇掉下平台则传回出生点"],
    ],
    widths=[1.35, 5.75],
    font_size=8.65,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "9.5 废铁甲虫", 2)
add_table(
    doc,
    ["项目", "规格"],
    [
        ["基础数值", "生命 35；冲撞伤害 12；感知距离 5.5；冲刺速度 11；完整冷却 1.8 秒"],
        ["锁定", "玩家在相同地面段且距离合适时进入蓄力；蓄力开始后锁定水平朝向"],
        ["蓄力", "0.80 秒；地面出现窄长警示带，甲壳闪动并播放摩擦声；受毒减速影响蓄力动画但总时长不低于 0.8 秒"],
        ["冲刺", "最大 4.8 单位或 0.45 秒；冲刺中免疫减速；接触玩家 Hurtbox 造成一次 12 伤害"],
        ["撞墙", "Rigidbody Cast 或碰撞法线确认前方 Ground 后立即停止，进入 Stunned 1.0 秒"],
        ["未撞墙", "到最大距离后进入 Recover 0.45 秒，不连续回头冲刺"],
        ["掉落", "Encounter03 的甲虫死亡必掉铁甲根"],
        ["防卡死", "冲刺不能越过 Encounter 边界；撞门视为撞墙；状态超时强制 Recover"],
    ],
    widths=[1.35, 5.75],
    font_size=8.55,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "9.6 敌人攻击节流", 2)
add_bullets(
    doc,
    [
        "EncounterDirector 提供最多 3 个同时攻击许可。进入 Telegraph 前申请许可，进入 Recover 或 Dead 时归还。",
        "未拿到许可的敌人可以追踪和重新站位，但不能开始前摇。第三战斗区最多 4 敌人，节流避免同帧无解攻击。",
        "攻击许可不影响已开始的攻击；暂停、嫁接和死亡时统一冻结或清空许可。",
        "若许可持有超过 4 秒仍未进入 Recover，EncounterDirector 强制回收并记录 Warning，避免异常状态锁死。",
    ],
)

add_heading(doc, "10 战斗区 掉落与关卡推进", 1)
add_body(doc, "每个 EncounterZone 独立拥有状态、门、出生点和存活敌人集合。LevelFlow 只监听区域 Cleared，不读取全局敌人列表。这样可以保证后续区域不会覆盖当前计数，也不会因为场外测试敌人阻止传送门开启。")

add_heading(doc, "10.1 Encounter 状态", 2)
add_table(
    doc,
    ["状态", "进入动作", "允许转换", "关键保护"],
    [
        ["Inactive", "门开放，敌人未生成或 Dormant", "Locked", "触发器只能由 Player Layer 激活"],
        ["Locked", "关闭左右门，锁定回退范围", "Active", "等待 0.35 秒关门反馈后生成敌人"],
        ["Active", "注册本区敌人并显示剩余数", "Cleared", "只接受自己注册敌人的死亡事件"],
        ["Cleared", "打开门，生成奖励，发出一次 Cleared", "无", "bool cleared 阻止重复结算"],
    ],
    widths=[1.0, 2.75, 1.25, 2.1],
    font_size=8.55,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "10.2 三个战斗区配置", 2)
add_table(
    doc,
    ["区域", "敌人配置", "教学目标", "必掉奖励", "清场结果"],
    [
        ["Encounter01", "藤蔓 x 2", "普攻 跳跃 冲刺 近战前摇", "藤蔓触须", "开右门，提示拾取和嫁接"],
        ["Encounter02", "藤蔓 x 2  蘑菇 x 1", "远程躲避并验证藤蔓攻击", "毒菌伞", "开右门，提示花部技能变化"],
        ["Encounter03", "藤蔓 x 2  蘑菇 x 1  甲虫 x 1", "综合战斗并验证毒雾和冲刺", "铁甲根", "清弹幕，通知 Portal 开启"],
    ],
    widths=[1.1, 1.6, 2.2, 1.2, 1.0],
    font_size=8.25,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "10.3 敌人注册与清场", 2)
add_code(
    doc,
    [
        "EncounterZone.StartEncounter",
        "  if state is not Inactive return",
        "  state becomes Locked",
        "  close gates",
        "  create enemies from EncounterDefinition",
        "  add each EnemyController to livingEnemies and subscribe Died",
        "  state becomes Active",
        "EncounterZone.OnEnemyDied enemy",
        "  if state is not Active return",
        "  if livingEnemies.Remove enemy is false return",
        "  if livingEnemies.Count equals 0 call CompleteOnce",
    ],
)

add_heading(doc, "10.4 必掉与拾取规则", 2)
add_numbered(
    doc,
    [
        "EncounterDefinition 配置 guaranteedGraft 和 guaranteedEnemyIndex。指定敌人死亡时生成一次奖励，不使用随机概率。",
        "GraftPickup 落地后保持存在，只有 Player Layer 可拾取。接触后先写入 GraftInventory，成功后再销毁自身。",
        "Inventory 以部件 ID 去重。重复拾取只刷新提示，不增加数量，也不影响已装备状态。",
        "若掉落物落入 KillPlane，则移动到本区域 RewardPoint；不得永久丢失。",
        "区域清场 2 秒后如果奖励尚未生成或引用丢失，EncounterZone 在 RewardPoint 补发一次并记录 Warning。",
        "玩家重开 Level01 时 GameSession.ResetRun 清空 Inventory 和 Loadout，三个默认槽位恢复基础能力。",
    ],
)

add_heading(doc, "10.5 门和区域边界", 2)
add_bullets(
    doc,
    [
        "Gate 关闭时同时启用 Collider 和关闭动画；动画只负责表现，Collider 是真实阻挡。",
        "区域开始后左门也关闭，防止玩家把敌人拉出战斗区。清场后左右门均开放。",
        "敌人的 AllowedBounds 来自所属 Encounter，Chase 和 Dash 每帧检查该边界。",
        "玩家不能在一个 Encounter 为 Active 时触发后续 Encounter。LevelFlow 额外验证前序区域全部 Cleared。",
    ],
)

add_heading(doc, "11 传送门 死亡与结算", 1)

add_heading(doc, "11.1 ExitPortal 状态", 2)
add_table(
    doc,
    ["状态", "表现", "Collider", "交互"],
    [
        ["Dormant", "枯死根冠或关闭外观", "交互关闭", "无提示"],
        ["Growing", "0.8 秒生长动画和粒子", "交互关闭", "忽略 E"],
        ["Open", "稳定循环光效", "交互 Trigger 开启", "范围内显示 E 进入"],
        ["Entering", "玩家被锁定并吸向中心 0.6 秒", "交互关闭", "重复输入无效"],
        ["Complete", "切换 Result 后不再更新", "关闭", "无"],
    ],
    widths=[1.0, 2.55, 1.55, 2.0],
    font_size=8.65,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "11.2 开门前后的程序顺序", 2)
add_numbered(
    doc,
    [
        "Encounter03 首次 Cleared 后，LevelFlow 标记 levelCleared，禁止新敌人和新攻击生成。",
        "ProjectileRegistry 销毁或失活所有 EnemyProjectile，敌人死亡协程可以继续播放但不再影响逻辑。",
        "ExitPortal.BeginOpen 从 Dormant 进入 Growing。0.8 秒后进入 Open 并启用交互 Trigger。",
        "玩家进入范围时 InteractionService 保存当前 Portal；离开范围立即清除提示。",
        "按 E 后 Portal 先设置 Entering，再请求 GameStateController 锁定 Gameplay，避免同帧重复调用。",
        "记录 runElapsed、三个槽位 ID 和完成标记，然后显示 ResultPanel。首版不写入磁盘。",
    ],
)

add_heading(doc, "11.3 玩家死亡", 2)
add_table(
    doc,
    ["时间点", "程序动作", "禁止发生"],
    [
        ["生命到 0 当帧", "Health 标记死亡  PlayerState 进入 Dead  关闭 Hitbox 与 Hurtbox", "再次扣血 冲刺 攻击 嫁接"],
        ["0 至 0.8 秒", "保留 timeScale 1 播放死亡动画和音效，敌人停止攻击意图", "敌人继续造成伤害或开启下一区域"],
        ["0.8 秒后", "GameState 进入 Dead，timeScale 0，显示死亡页", "自动复活"],
        ["选择重开", "timeScale 先恢复 1  ResetRun  重载 Level01", "保留库存 冷却 静态事件"],
        ["返回主页面", "timeScale 先恢复 1  清理 Session  加载 MainMenu", "旧音乐或场景对象继续存在"],
    ],
    widths=[1.25, 3.6, 2.25],
    font_size=8.45,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "11.4 结果页", 2)
add_bullets(
    doc,
    [
        "必须显示关卡完成、本局用时、根茎花三个最终装备，以及再次挑战和返回主页面。",
        "进入 Result 时 Time.timeScale 为 0，Gameplay Map 关闭，UI Map 开启。",
        "再次挑战执行完整场景重载，不手动逐个复位敌人。比赛版本优先可靠性。",
        "结果页只允许打开一次。任何晚到的敌人死亡、投射物或区域事件都必须被终止态过滤。",
    ],
)

add_heading(doc, "12 页面与 HUD 程序需求", 1)
add_body(doc, "正式界面全部使用 uGUI Canvas 和 TextMeshPro。CanvasScaler 设为 Scale With Screen Size，参考分辨率 1920 x 1080，Match 0.5。所有元素使用锚点布局，不允许以 Screen.width 和固定 Rect 手算位置。")

add_heading(doc, "12.1 主页面", 2)
add_table(
    doc,
    ["控件", "行为", "异常处理"],
    [
        ["开始游戏", "屏蔽自身后调用 SceneLoader.LoadLevel01", "加载失败恢复按钮并显示错误"],
        ["操作", "打开操作面板，返回后恢复原焦点", "不改变 GameState"],
        ["音乐音量", "写入 AudioMixer Music 参数与 PlayerPrefs", "初次运行默认 80%"],
        ["音效音量", "写入 AudioMixer SFX 参数与 PlayerPrefs", "初次运行默认 90%"],
        ["退出", "Editor 中停止播放，Build 中 Application.Quit", "按钮只执行一次"],
    ],
    widths=[1.35, 3.6, 2.15],
    font_size=8.65,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "12.2 HUD", 2)
add_table(
    doc,
    ["元素", "数据来源", "刷新方式", "要求"],
    [
        ["生命条", "PlayerHealth", "Damaged Healed Reset", "数值变化 0.15 秒插值，死亡归零"],
        ["三槽图标", "GraftLoadout", "GraftApplied Reset", "基础状态也有明确图标或空槽样式"],
        ["技能冷却", "PlayerCombat", "每帧只更新填充比例", "显示当前花部技能图标"],
        ["区域目标", "EncounterZone", "Started EnemyCountChanged Cleared", "显示清除敌人和剩余数量"],
        ["拾取提示", "GraftInventory", "ItemAdded", "显示部件名 槽位和按键提示约 3 秒"],
        ["交互提示", "InteractionService", "TargetChanged", "仅在传送门范围内出现"],
    ],
    widths=[1.2, 1.65, 2.15, 2.1],
    font_size=8.45,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "12.3 暂停 死亡和结果页面", 2)
add_table(
    doc,
    ["页面", "按钮", "打开条件", "关闭条件"],
    [
        ["PausePanel", "继续  音量  重开  返回主页面", "Playing 按 Esc", "继续或换场景"],
        ["DeadPanel", "重新开始  返回主页面", "死亡动画完成", "换场景"],
        ["ResultPanel", "再次挑战  返回主页面", "Portal Entering 完成", "换场景"],
    ],
    widths=[1.35, 2.5, 1.75, 1.5],
    font_size=8.75,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "12.4 UI 可靠性", 2)
add_bullets(
    doc,
    [
        "每个 Panel 只由对应 Presenter 控制显隐。其他系统通过事件传数据，不直接 FindObjectOfType 查 UI。",
        "按钮回调在 Inspector 或 OnEnable 中绑定一次，OnDisable 解除；禁止重复 AddListener。",
        "16 比 9 的 1280 x 720 和 1920 x 1080 必须无裁切；16 比 10 只要求不遮住按钮与关键信息。",
        "所有 UI 动画使用 unscaledDeltaTime，确保 timeScale 为 0 时仍能打开、选择和关闭。",
        "中文字体资产必须包含本项目所有字符。缺字在第一次导入时解决，不能临近提交动态生成大字库。",
    ],
)

add_heading(doc, "13 相机 动画 音频与美术接口", 1)

add_heading(doc, "13.1 相机", 2)
add_table(
    doc,
    ["功能", "实现", "参数与限制"],
    [
        ["跟随", "CameraFollow2D 在 LateUpdate 追踪玩家", "水平 damping 0.16 秒，垂直 damping 0.24 秒"],
        ["前视", "按玩家水平速度提供偏移", "最大 1.3 单位，换向时平滑过渡"],
        ["边界", "LevelCameraBounds 限制相机中心", "不显示关卡外空白区域"],
        ["震动", "CameraShakeService 叠加短时噪声", "普通命中 0.08 秒，甲虫撞墙 0.16 秒；暂停时停止"],
        ["传送门", "Entering 时短暂跟随 Portal 中心", "不做复杂过场，0.6 秒结束"],
    ],
    widths=[1.2, 3.4, 2.5],
    font_size=8.65,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "13.2 Player Animator 参数", 2)
add_table(
    doc,
    ["参数", "类型", "设置者", "用途"],
    [
        ["Speed", "float", "PlayerAnimationBridge", "Idle 和 Run 混合"],
        ["VerticalSpeed", "float", "PlayerAnimationBridge", "Jump 和 Fall"],
        ["Grounded", "bool", "PlayerMotor2D", "落地切换"],
        ["Attack", "trigger", "PlayerCombat", "默认或藤蔓普攻"],
        ["Skill", "trigger", "PlayerCombat", "种子或毒雾"],
        ["Dash", "trigger", "PlayerMotor2D", "冲刺"],
        ["Hurt", "trigger", "PlayerHealth", "受击"],
        ["Dead", "bool", "PlayerHealth", "死亡并保持"],
        ["StemType", "int", "GraftApplier", "0 基础  1 藤蔓"],
        ["FlowerType", "int", "GraftApplier", "0 种子  1 毒菌"],
        ["RootType", "int", "GraftApplier", "0 基础  1 铁甲"],
    ],
    widths=[1.35, 1.0, 2.25, 2.5],
    font_size=8.45,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "13.3 动画与判定接口", 2)
add_bullets(
    doc,
    [
        "攻击时序由 AttackDefinition 的 startup active recovery 驱动。Animator 负责表现，不能再额外用 Animation Event 开第二次 Hitbox。",
        "美术可以在动画中加入仅表现用事件，例如 PlayWhoosh 或 SpawnLeafVfx，但伤害生效点只有 AttackTimeline 一个来源。",
        "角色和敌人 Sprite 默认 Pivot 使用 Bottom Center。角色 RootVisualPoint StemVisualPoint FlowerVisualPoint 由美术在 Prefab 中确认位置。",
        "Collider 和 Hurtbox 不随攻击动画大幅缩放。需要变化的攻击范围使用独立 Hitbox Prefab。",
        "每个 Animator Controller 必须有无素材时可运行的占位状态，缺某段动画不会阻断逻辑。",
    ],
)

add_heading(doc, "13.4 音频程序接口", 2)
add_table(
    doc,
    ["事件 ID", "触发时机", "播放方式", "去重规则"],
    [
        ["ui_click", "按钮成功执行", "UI 总线 2D", "一次提交一次"],
        ["player_attack", "普攻前摇结束", "Player AudioSource", "一次攻击一次"],
        ["seed_shot", "生成种子投射物", "Player AudioSource", "一次技能一次"],
        ["poison_cast", "生成 PoisonZone", "Player AudioSource", "一次技能一次"],
        ["enemy_telegraph", "敌人进入 Telegraph", "敌人位置 3D Blend 0", "每个攻击一次"],
        ["hit_light", "有效 DamageInfo 被接受", "命中点 one shot", "多目标同帧最多 3 声"],
        ["player_hurt", "玩家有效受伤", "Player AudioSource", "受击无敌期内不重复"],
        ["graft_apply", "Graft 事务成功", "UI 与角色各一层", "失败不播放"],
        ["portal_open", "Portal Growing 开始", "Portal AudioSource", "关卡一次"],
        ["level_clear", "Portal 进入 Open", "Music 或 UI 总线", "关卡一次"],
    ],
    widths=[1.45, 2.7, 1.7, 1.25],
    font_size=8.3,
    alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "13.5 Layer 与 Sorting Layer", 2)
add_table(
    doc,
    ["名称", "类型", "允许交互", "备注"],
    [
        ["Ground", "Physics Layer", "Player Enemy Projectile", "TilemapCollider 与门"],
        ["PlayerBody", "Physics Layer", "Ground", "角色实体碰撞"],
        ["PlayerHurtbox", "Physics Layer", "EnemyHitbox EnemyProjectile", "只接收伤害"],
        ["EnemyBody", "Physics Layer", "Ground", "敌人实体可与玩家穿过，避免推挤"],
        ["EnemyHurtbox", "Physics Layer", "PlayerHitbox PlayerProjectile PoisonZone", "只接收伤害"],
        ["PlayerShield", "Physics Layer", "EnemyProjectile", "铁甲冲刺时启用"],
        ["Interactable", "Physics Layer", "Player 触发器", "掉落和 Portal"],
        ["Background Terrain Actor VFX Foreground UI", "Sorting Layer", "按顺序渲染", "危险预警位于 Terrain 与 Actor 之间或专设 Telegraph"],
    ],
    widths=[1.65, 1.35, 2.35, 1.75],
    font_size=8.25,
    alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "14 数据资产和脚本职责", 1, page_break=True)

add_heading(doc, "14.1 ScriptableObject 清单", 2)
add_table(
    doc,
    ["数据资产", "数量", "关键字段", "校验"],
    [
        ["PlayerConfig", "1", "生命 移动 跳跃 冲刺 无敌 受击硬直", "所有时间和速度大于 0"],
        ["AttackDefinition", "4", "abilityType damage startup active recovery cooldown range size projectile status vfx sfx", "时序非负；Prefab 与类型匹配"],
        ["EnemyDefinition", "3", "id health speed attackDamage ranges timings prefab guaranteedDrop", "ID 唯一；Prefab 有 EnemyController"],
        ["GraftDefinition", "3", "id slot icon visual ability attackData reduction blockProjectile description", "ID 唯一；槽位和能力一致"],
        ["EncounterDefinition", "3", "spawn entries gate delay guaranteedGraft guaranteedEnemyIndex", "敌人索引有效；第三关含甲虫"],
        ["AudioCue", "按事件", "clips volume pitchRange mixerGroup", "至少一个 Clip；音量 0 至 1"],
    ],
    widths=[1.45, 0.65, 3.6, 1.4],
    font_size=8.25,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "14.2 AttackDefinition 字段", 2)
add_table(
    doc,
    ["字段", "类型", "含义", "使用方"],
    [
        ["id", "string", "稳定 ID", "校验与调试"],
        ["executor", "AttackExecutorType", "MeleeBox VineLine Projectile PoisonZone", "PlayerCombat"],
        ["damage", "float", "直接伤害或每 Tick 伤害", "DamageInfo"],
        ["startup active recovery", "float", "攻击三个时段", "AttackTimeline"],
        ["cooldown", "float", "技能冷却或最短攻击间隔", "PlayerCombat"],
        ["offset size range", "Vector2 float", "近战和范围位置", "Hitbox Executor"],
        ["maxTargets", "int", "一次攻击最大不同目标", "Hitbox2D"],
        ["projectilePrefab speed lifetime", "Prefab float", "投射物参数", "Projectile Executor"],
        ["statusDuration tickInterval slow", "float", "毒状态参数", "PoisonZone StatusController"],
        ["vfxPrefab audioCue", "Reference", "表现事件", "VFX Audio 服务"],
    ],
    widths=[1.65, 1.45, 2.6, 1.4],
    font_size=8.2,
    alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "14.3 完整脚本清单", 2)
add_table(
    doc,
    ["脚本", "单一职责", "主要公开接口", "依赖"],
    [
        ["GameBootstrap", "创建并保留唯一的全局服务", "Initialize", "GameSession State Audio SceneLoader"],
        ["GameSession", "保存本局计时 库存 装备与结算快照", "BeginRun ResetRun CompleteRun", "无场景对象"],
        ["GameStateController", "管理状态 输入 Map 和 timeScale", "SetState CanTransition", "InputReader"],
        ["SceneLoader", "屏蔽重复加载并处理失败", "LoadMenu LoadLevel01 ReloadLevel", "SceneManager"],
        ["InputReader", "把 InputAction 转为语义事件", "Move Jump Attack Skill Dash Interact Graft Pause", "InputActionAsset"],
        ["PlayerMotor2D", "移动 跳跃 冲刺 GroundCheck", "SetInput BeginDash LockMovement", "Rigidbody2D PlayerConfig"],
        ["PlayerCombat", "攻击状态 冷却和执行器切换", "RequestAttack RequestSkill CancelAll ApplyLoadout", "AttackDefinition"],
        ["PlayerHealth", "生命 无敌 受击 死亡", "TryDamage ResetHealth", "Health Hurtbox"],
        ["PlayerAnimationBridge", "Animator 参数 外观挂点 受击闪色", "ApplyGraftVisual PlayFeedback", "Animator SpriteRenderer"],
        ["DamageInfo", "一次伤害请求的数据", "struct", "无"],
        ["Hitbox2D", "攻击窗口与目标去重", "Open Close", "ContactFilter2D"],
        ["Hurtbox2D", "接收并转交伤害", "Receive", "Health"],
        ["Projectile2D", "移动 碰撞 超时 阵营", "Launch Despawn", "DamageInfo"],
        ["PoisonZone", "挂载毒状态", "Activate", "StatusController"],
        ["StatusController", "毒伤和减速唯一实例", "ApplyPoison Clear", "EnemyController Health"],
        ["GraftInventory", "本局部件解锁与去重", "Add Contains Reset", "GameSession"],
        ["GraftLoadout", "根茎花当前装备", "Equip Get Reset", "GraftDefinition"],
        ["GraftApplier", "执行装备事务和通知", "TryApply", "Loadout Combat Animation HUD"],
        ["EnemyController", "敌人通用状态和死亡上报", "Activate ReturnToSpawn", "EnemyDefinition Encounter"],
        ["VineEnemy", "藤蔓追踪与近战序列", "TickBehaviour", "EnemyController EnemyAttack"],
        ["MushroomEnemy", "站位 瞄准与毒球", "TickBehaviour", "EnemyController Projectile"],
        ["BeetleEnemy", "蓄力 冲刺 撞墙眩晕", "TickBehaviour", "EnemyController Rigidbody2D"],
        ["EncounterZone", "本区状态 敌人集合 门和奖励", "Begin CompleteOnce", "EncounterDefinition Gate"],
        ["LevelFlow", "顺序验证和第三关完成处理", "RegisterEncounter OnEncounterCleared", "Portal ProjectileRegistry"],
        ["ExitPortal", "开门 交互和进入结算", "BeginOpen TryEnter", "State Session"],
        ["InteractionService", "维护当前最近交互对象", "SetTarget Clear Interact", "InputReader HUD"],
        ["MenuPresenter", "主页面按钮和设置", "Show Hide", "SceneLoader Audio"],
        ["HudPresenter", "订阅生命 装备 区域和提示", "Bind Unbind", "玩家和 LevelFlow"],
        ["GraftPresenter", "库存选择 对比 确认和取消", "Open Refresh Confirm Close", "GraftApplier State"],
        ["EndPresenter", "死亡与结果页面", "ShowDead ShowResult", "Session SceneLoader"],
        ["AudioService", "Mixer 音量 音乐与 OneShot", "Play SetVolume", "AudioCue"],
        ["CameraFollow2D", "跟随 前视 边界", "SetTarget SetBounds", "Player"],
    ],
    widths=[1.55, 2.35, 2.15, 1.05],
    font_size=7.65,
    alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "14.4 数据校验", 2)
add_bullets(
    doc,
    [
        "在 Editor 菜单 PlantSpirit Validate Demo Content 中扫描全部 Player Attack Enemy Graft Encounter 数据。",
        "阻断构建的错误包括空 ID、重复 ID、缺 Prefab、槽位与能力类型不符、Encounter 无敌人、必掉索引越界、第三关未绑定 Portal。",
        "运行时 Awake 再做轻量引用校验，错误日志必须包含场景对象路径或资产名，不输出无法定位的 NullReference。",
        "构建前自动确认 Build Settings 顺序为 MainMenu 和 Level01，且两个场景引用有效。",
    ],
)

add_heading(doc, "14.5 当前代码迁移表", 2)
add_table(
    doc,
    ["当前文件", "保留内容", "必须停止的内容", "迁移目标"],
    [
        ["PlantSpiritDemo.cs", "可参考旧数值", "整套顶视角原型和 OnGUI", "移出 Assets，不参与构建"],
        ["VerticalSliceRuntime.cs", "GroundCheck 和基础移动数值可参考", "自动启动 创建灰盒 单例列表 污染 融合 Respawn", "拆到 Core Player Combat"],
        ["Ggj48hFlow.cs", "三段区域顺序和必掉概念", "按 X 坐标直接推进 全局 activeEnemies OnGUI", "LevelFlow EncounterZone Canvas Presenter"],
        ["EnemiesAndBoss.cs", "三种敌人数值", "距离模拟碰撞 Boss 和额外敌人", "三个 Enemy Prefab 与行为脚本"],
        ["RunFlowAndUi.cs", "无 P0 运行逻辑需要保留", "随机路线 进化 存档 Boss HUD", "整体从首版程序集移除"],
        ["ContentDefinitions.cs", "部件 ID 和三项数值", "静态 Catalog Fusion Evolution 多余内容", "ScriptableObject 数据资产"],
        ["RunSeed.cs", "后续随机系统可参考", "首版无调用必要", "移到 Archive，第二阶段再启用"],
    ],
    widths=[1.45, 1.75, 2.35, 1.55],
    font_size=8.1,
    alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "15 程序任务拆分与 48 小时排期", 1)
add_body(doc, "以下排期以 1 名主程序为基准，给程序预留约 34 个有效开发小时，其余时间用于睡眠、沟通、素材接入和最终构建。若有 2 名程序，第二人优先负责敌人、Encounter 与测试，不要另开融合或 Boss 分支。")

add_heading(doc, "15.1 里程碑", 2)
add_table(
    doc,
    ["时间", "必须达到的可玩状态", "未达到时立即处理"],
    [
        ["H0 至 H2", "旧代码隔离  两场景可加载  Build Settings 正确", "停止所有内容制作，先修启动链"],
        ["H2 至 H8", "正式 Player Prefab 可移动 跳跃 冲刺，镜头和碰撞稳定", "减少移动动画，不减少控制功能"],
        ["H8 至 H14", "玩家可攻击测试假人，伤害 受击 死亡完整", "不做命中停顿美化，先修伤害去重"],
        ["H14 至 H20", "三种敌人可单独战斗，前摇和死亡正确", "减少动画差异，保留三种行为"],
        ["H20 至 H26", "三个 Encounter 可顺序清场，三个器官必掉", "取消额外波次，使用一次性生成"],
        ["H26 至 H31", "随时嫁接完成，三种效果真实可验", "界面先简化，但不得简化行为变化"],
        ["H31 至 H35", "传送门 死亡 结果 重玩 主页面闭环", "冻结新反馈，修所有终止态"],
        ["H35 至 H40", "接入正式美术 UI 音效，无引用断裂", "占位素材可保留，但流程不得回退"],
        ["H40 至 H44", "完成 15 条主验收，生成 RC1", "只修 P0，不改架构"],
        ["H44 至 H48", "第二台电脑验证 视频 截图 最终包", "保留已验证 RC，不覆盖唯一可用包"],
    ],
    widths=[1.0, 4.15, 1.95],
    font_size=8.25,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "15.2 可执行任务单", 2)
add_table(
    doc,
    ["ID", "任务", "估时", "依赖", "完成结果"],
    [
        ["PRG 01", "备份并隔离旧 Runtime 与 VerticalSlice", "0.5h", "无", "Play 不再自动生成灰盒"],
        ["PRG 02", "创建 MainMenu Level01 和 Build Settings", "0.8h", "01", "Build 从主页面启动"],
        ["PRG 03", "建立 asmdef Layers Sorting Layers Tags", "0.7h", "01", "碰撞矩阵和编译边界确定"],
        ["PRG 04", "创建 InputAction 和 InputReader", "1.0h", "03", "状态切换无重复输入"],
        ["PRG 05", "GameSession StateController SceneLoader", "1.5h", "02 04", "主页面可进关和返回"],
        ["PRG 06", "Player Prefab Rigidbody Collider GroundCheck", "1.0h", "03", "稳定站立和落地"],
        ["PRG 07", "移动 加减速 跳跃 缓冲 土狼时间", "2.0h", "06", "移动测试通过"],
        ["PRG 08", "冲刺 冷却 无敌 撞墙截短", "1.5h", "07", "冲刺不穿墙"],
        ["PRG 09", "Health DamageInfo Hitbox Hurtbox", "2.0h", "03", "伤害去重和死亡事件"],
        ["PRG 10", "PlayerCombat 默认普攻", "1.5h", "09", "假人命中一次"],
        ["PRG 11", "种子投射物与技能冷却", "1.0h", "09 10", "碰撞和超时正确"],
        ["PRG 12", "玩家受击 击退 无敌 死亡页触发", "1.2h", "09", "死亡不自动复活"],
        ["PRG 13", "EnemyController 通用状态和区域边界", "1.5h", "09", "Idle Chase Dead 可复用"],
        ["PRG 14", "藤蔓怪", "1.2h", "13", "前摇近战与掉落事件"],
        ["PRG 15", "蘑菇与敌方投射物", "1.4h", "13", "瞄准 发射 挡墙"],
        ["PRG 16", "甲虫蓄力 冲刺 撞墙眩晕", "1.8h", "13", "无越界和状态卡死"],
        ["PRG 17", "Encounter Gate Spawn LivingSet", "2.0h", "14 至 16", "单区清场只触发一次"],
        ["PRG 18", "三个区域配置和顺序验证", "1.2h", "17", "不能跳区"],
        ["PRG 19", "GraftPickup Inventory 必掉补发", "1.0h", "17", "三器官必得"],
        ["PRG 20", "GraftDefinition Loadout Applier", "1.4h", "19", "可装备和卸下"],
        ["PRG 21", "藤蔓执行器", "0.9h", "20", "前方两目标穿透"],
        ["PRG 22", "毒雾与状态控制", "1.3h", "20", "DOT 不叠层且减速"],
        ["PRG 23", "铁甲减伤与冲刺护盾", "1.0h", "08 20", "12 伤变 9 并可挡弹"],
        ["PRG 24", "GraftPanel 正式 UI", "1.5h", "20", "战斗中暂停 选择 确认 取消"],
        ["PRG 25", "HUD 提示与三槽显示", "1.0h", "05 20", "事件刷新无轮询查找"],
        ["PRG 26", "Portal Interaction Result", "1.2h", "18", "第三区清场后完成"],
        ["PRG 27", "Pause Dead Result Menu 完整按钮", "1.0h", "05 12 26", "重开和返回无残留"],
        ["PRG 28", "相机 边界 轻震动", "0.8h", "07", "不显示关卡外区域"],
        ["PRG 29", "Animator 与挂点接口", "1.0h", "06 20", "美术替换不改代码"],
        ["PRG 30", "AudioMixer 与事件钩子", "0.8h", "05", "音量保存且不叠播"],
        ["PRG 31", "内容校验菜单和 Debug Overlay", "0.8h", "数据完成", "一键发现空引用"],
        ["PRG 32", "PlayMode 主流程与三次重玩测试", "1.5h", "全部 P0", "无软锁和残留"],
        ["PRG 33", "Windows 构建 第二机验证 提交包", "1.5h", "32", "无 Unity 环境通关"],
    ],
    widths=[0.65, 2.7, 0.65, 1.2, 1.9],
    font_size=7.65,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "15.3 两名程序分工", 2)
add_table(
    doc,
    ["角色", "主责", "不得同时修改", "每日合并点"],
    [
        ["程序 A", "Core Input Player Combat Grafting UI", "Enemy Prefab 与 EncounterDefinition", "H12 H24 H36"],
        ["程序 B", "Enemies Encounter Portal LevelFlow Tests", "Player Prefab 与 GameStateController", "H12 H24 H36"],
        ["共同", "数据字段命名 Layer 动画参数和发布构建", "未经沟通修改公共接口", "每次合并后立即跑完整流程"],
    ],
    widths=[1.0, 2.75, 2.35, 1.0],
    font_size=8.55,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
)

add_heading(doc, "16 测试 调试与验收", 1)

add_heading(doc, "16.1 必须提供的调试能力", 2)
add_table(
    doc,
    ["工具", "Development Build 操作", "用途", "Release Build"],
    [
        ["Debug Overlay", "F1", "显示 GameState Encounter 状态 玩家状态 冷却 敌人数", "关闭"],
        ["Grant Grafts", "F2", "把三器官加入库存，不自动装备", "关闭"],
        ["Kill Encounter", "F3", "对当前区域敌人造成致死伤害，走正常死亡事件", "关闭"],
        ["Toggle Invincible", "F4", "用于敌人和流程测试", "关闭"],
        ["Reload Level", "F5", "验证场景重载和静态清理", "关闭"],
    ],
    widths=[1.35, 1.45, 3.1, 1.2],
    font_size=8.55,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
)

add_heading(doc, "16.2 自动化测试最小集", 2)
add_table(
    doc,
    ["测试", "类型", "断言"],
    [
        ["ContentValidationTests", "EditMode", "3 个 Graft 3 个 Enemy 3 个 Encounter 的 ID 与引用全部有效"],
        ["DamageDeduplicationTests", "EditMode", "同一 attackInstanceId 对同一 Hurtbox 只扣一次"],
        ["GraftModifierTests", "EditMode", "铁甲根将 12 变 9；卸下恢复 12；毒状态刷新不叠 DPS"],
        ["EncounterCompletionTests", "PlayMode", "全部注册敌人死亡后只发一次 Cleared，其他区域敌人不影响"],
        ["GameStateTests", "PlayMode", "Playing Grafting Playing 时 timeScale 与 Input Map 正确"],
        ["RestartTests", "PlayMode", "重载后库存为空 装备基础 timeScale 1 静态事件无重复"],
    ],
    widths=[1.8, 1.0, 4.3],
    font_size=8.55,
    alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "16.3 主流程人工验收", 2)
add_table(
    doc,
    ["编号", "操作", "通过结果"],
    [
        ["AT 01", "从 Build 启动并点击开始", "进入 Level01，玩家和 HUD 仅有一个实例"],
        ["AT 02", "测试移动 跳跃 缓冲 土狼时间 冲刺", "无穿墙 无贴墙误判 无失控滑动"],
        ["AT 03", "同一次普攻持续接触敌人", "只造成一次 10 伤害"],
        ["AT 04", "进入 Encounter01 后尝试前进和后退", "门阻挡，不能触发 Encounter02"],
        ["AT 05", "击杀 Encounter01 首只藤蔓", "必掉藤蔓触须且区域计数正确"],
        ["AT 06", "敌人攻击中打开嫁接", "玩家 敌人 投射物 DOT 全部暂停，UI 可操作"],
        ["AT 07", "装备藤蔓触须并攻击前方两敌", "2.8 距离内最多两个目标各受 8 伤害"],
        ["AT 08", "装备毒菌伞并施放技能", "无种子生成，毒雾 3 秒 6 次结算并减速 30%"],
        ["AT 09", "重复用毒雾覆盖同一敌人", "持续时间刷新，DPS 不翻倍"],
        ["AT 10", "装备铁甲根承受甲虫 12 伤害", "实际扣 9；卸下后实际扣 12"],
        ["AT 11", "铁甲冲刺前段撞上蘑菇毒球", "毒球销毁并播放挡弹；后段接触仍受伤"],
        ["AT 12", "清除前两区并触发第三战斗区", "区域顺序正确，前区敌人或尸体不影响计数"],
        ["AT 13", "击杀第三区域最后一敌", "敌方弹幕清理，0.8 秒后 Portal Open"],
        ["AT 14", "Portal 未开时按 E，开后连续按 E", "未开无效；开后只进入一次 Result"],
        ["AT 15", "在战斗中死亡", "0.8 秒后死亡页，无自动复活，无继续受伤"],
        ["AT 16", "死亡后重开", "库存 装备 敌人 区域 冷却全部恢复初始"],
        ["AT 17", "通关后再次挑战", "新局从教学起点开始，计时归零"],
        ["AT 18", "返回主页面后再开始", "音乐不叠播，按钮不重复触发"],
        ["AT 19", "1280 x 720 与 1920 x 1080 运行", "HUD 嫁接 死亡和结果按钮无裁切"],
        ["AT 20", "在第二台无 Unity 电脑连续通关 3 次", "无崩溃 无 Error 无软锁"],
    ],
    widths=[0.65, 3.15, 3.3],
    font_size=8.15,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "16.4 关键边界测试", 2)
add_bullets(
    doc,
    [
        "敌人死亡与玩家死亡发生在同一帧时，以玩家死亡为终止态，不自动打开 Portal；重开后恢复正常。",
        "玩家在 Portal 范围内打开嫁接时，关闭嫁接不能自动进入 Portal，必须重新按 E。",
        "敌方投射物在打开嫁接前一帧接近玩家，暂停期间不能继续移动或延迟命中。",
        "甲虫冲刺时玩家越过其身后，甲虫保持锁定方向直到冲刺结束，不能瞬间回头。",
        "掉落物生成位置低于 KillPlane 或被门挤出时，2 秒内回到 RewardPoint。",
        "快速连续按开始 重开 返回按钮时，SceneLoader 只接受第一个请求。",
    ],
)

add_heading(doc, "17 构建发布与删减规则", 1)

add_heading(doc, "17.1 构建设置", 2)
add_table(
    doc,
    ["项目", "要求"],
    [
        ["平台", "Windows x86 64"],
        ["场景顺序", "0 MainMenu  1 Level01"],
        ["分辨率", "默认 1920 x 1080，可窗口化；最小验证 1280 x 720"],
        ["帧率", "Application.targetFrameRate 60；VSync 由最终设备测试决定"],
        ["日志", "RC 使用 Development Build 测试；提交包使用非 Development Build"],
        ["目录", "Build Windows PlantSpirit.exe 与 PlantSpirit_Data 保持完整，不提交 Library Temp Logs"],
        ["说明", "README 写控制方式、运行环境、已知问题；credits 写外部素材和许可证"],
    ],
    widths=[1.5, 5.6],
    font_size=8.8,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "17.2 删减顺序", 2)
add_table(
    doc,
    ["顺序", "可删内容", "不能删的对应核心"],
    [
        ["1", "手柄绑定 复杂震屏 多音效随机", "键鼠 主要命中声 基础镜头"],
        ["2", "菜单动画 嫁接预览动画 通关过场", "页面可操作 状态正确"],
        ["3", "敌人受击动画差异和多段死亡粒子", "前摇 命中 死亡反馈"],
        ["4", "对象池 自动化测试中的非核心用例", "投射物超时 人工主流程验收"],
        ["5", "第三战斗区的一只藤蔓", "三种敌人和铁甲根掉落"],
    ],
    widths=[0.8, 3.15, 3.15],
    font_size=8.7,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)
add_body(doc, "任何情况下都不能删除随时嫁接、三种行为变化、区域独立清场、传送门完成、死亡重开和 Windows Build。它们共同构成评委实际体验到的首版闭环。")

add_heading(doc, "17.3 提交前检查", 2)
add_table(
    doc,
    ["确认项", "必须为真"],
    [
        ["编译", "Console 0 Error；正式程序集不引用 Archive 代码"],
        ["启动", "Build 默认进入 MainMenu，不依赖编辑器当前场景"],
        ["范围", "菜单和游戏内没有 Boss 融合 污染 进化 随机房间的残留入口"],
        ["流程", "三个区域按序完成，三器官必掉，Portal 只在最终清场后开启"],
        ["状态", "暂停 嫁接 死亡 结果的 timeScale 与 Input Map 全部正确"],
        ["重玩", "连续 3 次无静态引用 事件订阅 音乐 库存 冷却残留"],
        ["显示", "两种目标分辨率无 UI 裁切，中文无缺字"],
        ["外部电脑", "无需 Unity 和管理员权限即可启动并通关"],
        ["备份", "保留最后一个通过验证的 RC 压缩包，不被临时版本覆盖"],
    ],
    widths=[1.3, 5.8],
    font_size=8.8,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "18 程序开工清单", 1)
add_body(doc, "程序开始编码前只确认以下决定。确认后以本清单锁定首版范围，所有新想法进入赛后列表。")
add_table(
    doc,
    ["确认项", "最终决定"],
    [
        ["引擎", "Unity 2022.3.62f3，不升级"],
        ["场景", "MainMenu 和 Level01 两个正式场景"],
        ["输入", "Input System；键鼠 P0，手柄 P1"],
        ["角色", "动态 Rigidbody2D；移动 跳跃 冲刺 普攻 技能 受击 死亡"],
        ["嫁接", "战斗中可打开并暂停；根茎花三槽；无消耗，可反复替换"],
        ["部件", "藤蔓触须改普攻  毒菌伞改技能  铁甲根改减伤和冲刺挡弹"],
        ["敌人", "藤蔓怪 蘑菇 甲虫三种，均有可读前摇"],
        ["关卡", "教学起点 三个独立 Encounter 传送门出口"],
        ["掉落", "指定敌人必掉，不使用随机掉率"],
        ["通关", "第三 Encounter 清场后 Portal 开启，按 E 进入 Result"],
        ["死亡", "死亡页手动重开，不自动复活"],
        ["首版不做", "Boss 融合 污染 进化 随机房间 永久成长 单局存档"],
        ["冻结", "H38 内容冻结，H40 后只修 P0"],
        ["最终验收", "第二台电脑连续通关 3 次，20 条主流程验收全部通过"],
    ],
    widths=[1.35, 5.75],
    font_size=8.8,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
