from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"E:\26翌光游戏开发\策划文档\植物精灵 GGJ48H可玩Demo完整制作策划案 V1.0.docx")

FONT_CN = "Microsoft YaHei"
FONT_EN = "Aptos"
BLACK = "000000"
CHARCOAL = "36413B"
PALE = "F2F6F3"
LIGHT_GRAY = "D9D9D9"
MID_GRAY = "666666"
ACCENT = "2F6B45"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=110, bottom=100, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), LIGHT_GRAY)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_run_font(run, size=None, bold=None, color=BLACK, italic=None):
    run.font.name = FONT_EN
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_CN)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, size=10.5, color=BLACK, bold=False):
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.16
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold if bold else None)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=10.5, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=10.5)
    else:
        run = p.add_run(text)
        set_run_font(run, size=10.5)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.16
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=10.3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12


def add_numbered(doc, items):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.keep_together = True
        run = p.add_run(f"{index}.  {item}")
        set_run_font(run, size=10.3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12


def add_table(doc, headers, rows, widths=None, font_size=8.9, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    set_row_cant_split(hdr)
    for index, text in enumerate(headers):
        cell = hdr.cells[index]
        set_cell_shading(cell, CHARCOAL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        p.paragraph_format.keep_with_next = True
        run = p.add_run(str(text))
        set_run_font(run, size=font_size, bold=True, color="FFFFFF")
    for row_index, row_values in enumerate(rows):
        row = table.add_row()
        set_row_cant_split(row)
        for col_index, value in enumerate(row_values):
            cell = row.cells[col_index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index % 2 == 1:
                set_cell_shading(cell, PALE)
            p = cell.paragraphs[0]
            alignment = WD_ALIGN_PARAGRAPH.LEFT
            if alignments and col_index < len(alignments):
                alignment = alignments[col_index]
            p.alignment = alignment
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            run = p.add_run(str(value))
            set_run_font(run, size=font_size)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    keep_with_next(p)
    return p


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run, size=8.5, color=MID_GRAY)


_page_break_count = 0


def add_page_break(doc):
    # Keep only the cover break. Later sections flow naturally so a split table
    # cannot strand one row on an otherwise empty page.
    global _page_break_count
    if _page_break_count == 0:
        doc.add_page_break()
    _page_break_count += 1


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.62)
section.bottom_margin = Inches(0.62)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT_EN
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.16

title_style = styles["Title"]
title_style.font.name = FONT_EN
title_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
title_style.font.size = Pt(27)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor.from_string(BLACK)
title_ppr = title_style._element.get_or_add_pPr()
title_border = title_ppr.find(qn("w:pBdr"))
if title_border is not None:
    title_ppr.remove(title_border)

for style_name, size, before, after in (
    ("Heading 1", 17, 13, 6),
    ("Heading 2", 13, 10, 4),
    ("Heading 3", 11.2, 7, 3),
):
    style = styles[style_name]
    style.font.name = FONT_EN
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(BLACK)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_style_name in ("List Bullet", "List Number"):
    style = styles[list_style_name]
    style.font.name = FONT_EN
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    style.font.size = Pt(10.3)

# Footer
footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run("植物精灵 GGJ 48H 可玩 Demo 制作策划案  V1.0   |   ")
set_run_font(run, size=8.5, color=MID_GRAY)
add_page_field(footer_p)

# Cover
spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(36)
title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("植物精灵 GGJ 48H 可玩 Demo 制作策划案")
title_ppr = title._p.get_or_add_pPr()
title_border = title_ppr.find(qn("w:pBdr"))
if title_border is not None:
    title_ppr.remove(title_border)
for run in title.runs:
    set_run_font(run, size=27, bold=True, color=BLACK)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(22)
run = subtitle.add_run("版本 1.0  核心闭环执行版")
set_run_font(run, size=14, bold=True, color=BLACK)

intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
intro.paragraph_format.left_indent = Inches(0.55)
intro.paragraph_format.right_indent = Inches(0.55)
intro.paragraph_format.space_after = Pt(20)
intro.paragraph_format.line_spacing = 1.25
run = intro.add_run(
    "48 小时内交付一条 5 至 8 分钟、从主页面到第一关结算完整可玩的横版动作流程。"
    "玩家击败敌怪取得器官部件，可在战斗之外或战斗进行中随时打开嫁接界面，立即改变普攻、技能或冲刺表现；"
    "清理第一关全部敌人后，传送门生长并开启，进入后完成 Demo。"
)
set_run_font(run, size=11.5)

add_table(
    doc,
    ["项目", "首版决定"],
    [
        ["平台与引擎", "Windows PC，Unity 2022.3 LTS，16 比 9 横版"],
        ["单局时长", "首次游玩 5 至 8 分钟，熟练后 3 至 5 分钟"],
        ["核心卖点", "敌怪部件可随时嫁接，身体变化立刻改变操作结果"],
        ["第一关", "1 张手工关卡，3 个战斗区，3 种敌人，3 种嫁接部件"],
        ["通关方式", "全敌人清除后传送门出现并可交互"],
        ["明确延期", "融合台、融合配方、Boss、随机房间、污染、永久成长"],
    ],
    widths=[1.45, 5.65],
    font_size=9.5,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(14)
run = meta.add_run("更新日期 2026 年 9 月 4 日")
set_run_font(run, size=9.5, color=MID_GRAY)

add_page_break(doc)

add_heading(doc, "1 项目结论与范围", 1)
add_body(
    doc,
    "这版 Demo 的目标不是证明内容很多，而是让评委在几分钟内理解并亲手验证一句话：击败敌怪，拿走它的器官，"
    "把器官嫁接到自己身上，攻击方式马上发生变化。围绕这句话，主页面、第一关、玩家动作、三种敌人、随时嫁接、"
    "三类能力变化、清场传送门和结算页构成不可拆分的完整闭环。",
)

add_heading(doc, "1.1 嫁接与融合的边界", 2)
add_table(
    doc,
    ["系统", "首版规则", "频率与地点", "本次状态"],
    [
        ["嫁接", "把一个敌怪部件装入根 茎 花对应槽位，立即替换该槽位的能力与外观", "随时打开嫁接界面，战斗暂停后操作", "P0 必做"],
        ["融合", "把两个已嫁接特征合成为新的复合形态，拥有专属行为与代价", "只能在后续版本的融合台进行", "本次不做，仅预留数据扩展点"],
    ],
    widths=[0.85, 3.2, 1.65, 1.4],
    font_size=8.8,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
)
add_body(
    doc,
    "关键原则：嫁接是本作高频、即时、帮助玩家理解身体构筑的操作；融合是低频、高价值、需要配方与专属资产的深化系统。"
    "首版不出现融合按钮、融合提示或不可用入口，避免评委误以为功能损坏。程序只需保证特征数据以后可以增加复合效果。",
)

add_heading(doc, "1.2 功能优先级", 2)
add_table(
    doc,
    ["优先级", "定义", "内容"],
    [
        ["P0", "提交版本缺少任意一项即不成立", "主页面；第一关；移动跳跃冲刺；普攻与技能；3 种敌人；3 个嫁接部件；随时嫁接；清场传送门；死亡与通关结算；基础音画反馈"],
        ["P1", "P0 稳定后再加入", "暂停页；按键说明；轻量镜头震动；伤害数字；敌人剩余数；屏幕设置；手柄基础映射"],
        ["P2", "时间富余才加入", "二段跳；复杂连击；随机掉率；小地图；剧情对话；隐藏收集；更多敌人和部件"],
        ["后续", "本届提交后开发", "融合台与配方；Boss；根门多路线；随机房间；进化与污染；图鉴和永久解锁"],
    ],
    widths=[0.7, 1.65, 4.75],
    font_size=8.8,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "1.3 成功标准", 2)
add_bullets(
    doc,
    [
        "未听讲解的玩家能从主页面开始，在 8 分钟内完成第一关并进入传送门。",
        "玩家至少完成两次嫁接，并能用一句话说出嫁接前后普攻或技能的区别。",
        "任意嫁接都同时产生数值 行为 外观或特效中的至少两类变化，不能只改伤害数字。",
        "敌人攻击有清晰前摇，玩家受击 命中 清场 嫁接和传送门开启均有声音与画面反馈。",
        "Windows 构建可脱离 Unity 编辑器启动，完整流程无阻断性错误，目标为 1080p 60 FPS。",
    ],
)

add_page_break(doc)
add_heading(doc, "2 玩家流程与页面", 1)

add_heading(doc, "2.1 单局流程", 2)
add_table(
    doc,
    ["步骤", "玩家行为", "系统反馈", "目标时长"],
    [
        ["主页面", "点击开始游戏", "短转场进入枯萎林地入口", "10 至 20 秒"],
        ["基础教学", "移动 跳跃 普攻 技能 冲刺", "地面按键图标和一次性短提示", "30 至 45 秒"],
        ["战斗区一", "击败近战敌人并拾取藤蔓触须", "强制首掉，提示打开嫁接界面", "45 至 60 秒"],
        ["战斗区二", "应对远程敌人并拾取毒菌伞", "嫁接花部后技能变为毒雾", "60 至 90 秒"],
        ["战斗区三", "混合敌群与冲刺甲虫", "拾取铁甲根，可调整三槽构筑", "90 至 150 秒"],
        ["清场", "确认无存活敌人", "镜头轻推向出口，传送门生长并点亮", "5 至 10 秒"],
        ["通关", "走到门前按交互键", "进入 Demo 完成页，可重玩或回主菜单", "15 至 30 秒"],
    ],
    widths=[1.0, 2.3, 2.8, 1.0],
    font_size=8.7,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
)

add_heading(doc, "2.2 主页面", 2)
add_body(doc, "第一屏直接显示游戏名 植物精灵、枯萎森林背景与主角剪影，确保作品和主题在进入游戏前就可识别。")
add_table(
    doc,
    ["元素", "行为", "P0 验收"],
    [
        ["开始游戏", "默认选中；点击或确认后加载第一关", "连续点击不会重复加载"],
        ["操作说明", "打开一页简明键位图", "可关闭并回到主页面"],
        ["退出游戏", "Windows 构建退出；编辑器中不报错", "仅 PC 构建显示"],
        ["音量", "主音量滑杆或静音切换", "至少控制全部声音"],
    ],
    widths=[1.25, 3.05, 2.8],
    font_size=9.1,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "2.3 游戏内页面", 2)
add_table(
    doc,
    ["界面", "必须显示", "交互规则"],
    [
        ["HUD", "生命条；技能冷却；根 茎 花三个槽位；交互提示", "不遮挡脚下平台和敌方预警"],
        ["嫁接界面", "当前槽位；已收集部件；效果变化；确认与返回", "任何可操作时刻可打开；打开后暂停战斗"],
        ["暂停页", "继续；重新开始；返回主页面", "暂停音频与游戏时间"],
        ["死亡页", "重新挑战；返回主页面", "不需要永久成长或存档"],
        ["完成页", "本局用时；已嫁接部件；再次挑战；返回主页面", "进入传送门后 1 秒内出现"],
    ],
    widths=[1.15, 3.1, 2.85],
    font_size=9,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_page_break(doc)
add_heading(doc, "3 第一关枯萎林地入口", 1)
add_body(
    doc,
    "第一关使用一张手工制作的横版场景，不做程序随机。场景由教学起点、三个战斗区和出口组成，沿水平方向推进，"
    "仅加入低风险高差，避免平台跳跃抢走嫁接与战斗的注意力。战斗区之间用短走廊和单向封锁分隔，玩家可以在每次战斗后整理嫁接。",
)

add_heading(doc, "3.1 空间结构", 2)
add_table(
    doc,
    ["区域", "空间用途", "内容", "出口条件"],
    [
        ["起点", "教学和视觉建立", "平地 小台阶 木桩 静态污染；无敌人", "玩家越过教学标记"],
        ["战斗区一", "学习近战与首次嫁接", "宽 18 至 22 格；一处低平台；2 只藤蔓怪", "敌人全灭并拾取藤蔓触须"],
        ["战斗区二", "学习处理远程火力", "宽 22 至 26 格；上下两层；2 藤蔓怪加 1 蘑菇", "敌人全灭"],
        ["战斗区三", "混合能力检验", "宽 26 至 32 格；两处平台；2 藤蔓怪 1 蘑菇 1 甲虫", "关卡内全部敌人计数为零"],
        ["出口", "清场回报和结算", "枯死根冠在清场后生长为传送门", "玩家进入范围并按 E"],
    ],
    widths=[1.0, 1.65, 3.2, 1.25],
    font_size=8.65,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "3.2 战斗编排", 2)
add_table(
    doc,
    ["战斗区", "敌人配置", "部件保障", "教学目的"],
    [
        ["一", "藤蔓怪 2", "首只藤蔓怪必掉藤蔓触须", "理解命中 受击 拾取和茎部嫁接"],
        ["二", "藤蔓怪 2 蘑菇 1", "首只蘑菇必掉毒菌伞", "先处理远程单位，验证技能替换"],
        ["三", "藤蔓怪 2 蘑菇 1 甲虫 1", "甲虫必掉铁甲根", "利用范围 技能和冲刺对抗混编"],
    ],
    widths=[0.9, 1.65, 2.25, 2.3],
    font_size=9,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)
add_body(
    doc,
    "首版不要使用随机掉率。三个关键部件采用首杀必掉，保证每位评委都能看到核心系统。敌怪后续重复掉落只生成少量治疗花粉，"
    "避免地面出现无用重复部件。若玩家未拾取关键部件，战斗区出口旁生成一次补发拾取物。",
)

add_heading(doc, "3.3 传送门规则", 2)
add_numbered(
    doc,
    [
        "关卡载入时出口只显示枯死根冠，不可交互，碰撞不阻挡玩家。",
        "EncounterManager 注册所有战斗区敌人。只有已触发的战斗全部结束且存活敌人数为零，才发送 LevelCleared 事件。",
        "收到事件后等待 0.8 秒，播放根系生长动画、粒子和开启音效；HUD 显示 前往传送门。",
        "传送门开启后保持存在。玩家进入交互范围时显示 E 进入，不允许敌人死亡动画或投射物阻塞完成条件。",
        "交互后锁定输入，播放 0.6 秒吸入或淡出，切换到完成页。当前版本不实际加载第二关，避免出现空场景。",
    ],
)

add_page_break(doc)
add_heading(doc, "4 玩家角色与战斗", 1)

add_heading(doc, "4.1 操作映射", 2)
add_table(
    doc,
    ["动作", "键鼠", "手柄建议", "说明"],
    [
        ["移动", "A D 或方向键", "左摇杆", "有加减速，空中可修正"],
        ["跳跃", "Space", "A", "包含土狼时间和输入缓冲"],
        ["冲刺", "Left Shift", "RB", "短距离位移，带短暂无敌"],
        ["普攻", "鼠标左键或 J", "X", "朝角色面向方向攻击"],
        ["技能", "鼠标右键或 K", "Y", "默认种子射击，花部嫁接后替换"],
        ["交互", "E", "B", "拾取可自动，传送门需确认"],
        ["嫁接", "Tab 或 G", "LT", "随时打开嫁接界面并暂停战斗"],
        ["暂停", "Esc", "Menu", "继续 重开 返回主页面"],
    ],
    widths=[1.0, 1.6, 1.4, 3.1],
    font_size=8.8,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "4.2 基础数值", 2)
add_table(
    doc,
    ["参数", "推荐值", "验收意图"],
    [
        ["最大生命", "100", "允许承受约 8 至 10 次普通攻击"],
        ["地面速度", "5.2 单位每秒", "一屏移动不迟缓"],
        ["跳跃初速度", "12.5", "可越过约 3 格高差"],
        ["土狼时间", "0.12 秒", "平台边缘操作更宽容"],
        ["跳跃缓冲", "0.12 秒", "落地前按键仍能起跳"],
        ["冲刺距离", "3.6 单位", "跨越敌人与预警区"],
        ["冲刺时长", "0.22 秒", "方向明确，不做长距离飞行"],
        ["冲刺冷却", "1.1 秒", "可频繁使用但不能无脑连续闪避"],
        ["冲刺无敌", "0.18 秒", "覆盖冲刺前段"],
        ["受击无敌", "0.75 秒", "避免多敌重叠瞬杀"],
    ],
    widths=[1.6, 1.6, 3.9],
    font_size=9.1,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "4.3 默认攻击", 2)
add_table(
    doc,
    ["招式", "伤害与节奏", "判定", "表现要求"],
    [
        ["嫩芽抽打", "10 伤害；0.45 秒间隔；前摇 0.12 秒；后摇 0.18 秒", "前方 1.2 单位矩形；仅命中一次", "绿色弧形拖尾；轻命中停顿 0.04 秒；小怪击退"],
        ["种子射击", "8 伤害；冷却 3 秒", "直线投射物；速度 8.5；存在 1.2 秒", "发射闪光；飞行尾迹；命中小粒子"],
        ["基础冲刺", "无伤害；冷却 1.1 秒", "角色短暂无敌，不穿越墙体", "叶片残影；镜头轻微跟随加速"],
    ],
    widths=[1.1, 2.1, 1.8, 2.1],
    font_size=8.7,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)
add_body(
    doc,
    "首版只需一段稳定普攻，不做三段连击。核心展示来自嫁接后攻击形态的切换。若在第 30 小时前全部 P0 已稳定，可把第三次连续普攻做成重抽，"
    "否则不要增加动画和状态机分支。",
)

add_heading(doc, "4.4 伤害与反馈规则", 2)
add_bullets(
    doc,
    [
        "伤害判定与受击判定分离。攻击在动画事件点生成一次命中，不用每帧重复判定。",
        "命中敌人时必须同时出现受击闪白或染色、轻微击退、命中特效和命中音效。",
        "玩家受击时扣血、闪红、短暂无敌并受到小幅击退；无敌期间忽略后续伤害但仍可显示擦碰特效。",
        "敌方攻击前摇必须至少 0.45 秒，并使用姿态、地面形状或闪烁提示，不能只靠颜色。",
        "玩家死亡后禁用输入和伤害，播放死亡动画，再显示死亡页；不允许重复弹出结算。",
    ],
)

add_page_break(doc)
add_heading(doc, "5 随时嫁接系统", 1)
add_body(
    doc,
    "嫁接是首版唯一构筑系统。玩家拾取敌怪器官后，部件进入局内库存。按 Tab 或 G 可在任何时刻打开嫁接界面，"
    "游戏时间暂停。玩家把部件装入对应槽位并确认，能力与外观立即刷新。为了鼓励展示和测试，首版嫁接不消耗货币，旧部件返回库存，可反复切换。",
)

add_heading(doc, "5.1 三槽规则", 2)
add_table(
    doc,
    ["槽位", "职责", "默认状态", "首版部件"],
    [
        ["根", "冲刺和生存", "基础根系，无额外效果", "铁甲根"],
        ["茎", "普攻形态", "嫩芽抽打", "藤蔓触须"],
        ["花", "主动技能", "种子射击", "毒菌伞"],
    ],
    widths=[0.8, 1.7, 2.3, 2.3],
    font_size=9.2,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
)

add_heading(doc, "5.2 部件与能力变化", 2)
add_table(
    doc,
    ["部件", "来源与槽位", "实际效果", "可见变化", "音频反馈"],
    [
        ["藤蔓触须", "藤蔓怪 茎", "普攻变为 2.8 单位藤鞭；8 伤害；0.7 秒间隔；可穿透 1 个敌人", "手臂或茎挂点长出藤蔓；攻击轨迹更长；命中带叶片碎屑", "更长的挥鞭声；穿透命中增加清脆层"],
        ["毒菌伞", "蘑菇 花", "技能变为前方 3.0 单位毒雾；每秒 6 伤害；持续 3 秒；减速 30%；冷却 5 秒", "头部或花挂点出现菌伞；紫绿雾区边界清楚；敌人显示中毒图标", "喷雾声；持续轻气泡声；中毒命中弱提示"],
        ["铁甲根", "甲虫 根", "受伤减免 25%；冲刺前 0.18 秒获得树皮护盾并可撞碎敌方普通投射物", "脚部根系变粗并带铁灰树皮；冲刺出现护盾轮廓和碎屑", "低频木甲声；挡弹时独立脆裂声"],
    ],
    widths=[1.05, 1.05, 2.45, 1.75, 0.8],
    font_size=8.2,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "5.3 嫁接交互流程", 2)
add_numbered(
    doc,
    [
        "玩家接触部件后自动拾取，屏幕右侧显示部件名 槽位 核心效果，并提示 Tab 嫁接。",
        "打开界面后 Time.timeScale 设为零，敌人 投射物和伤害暂停，UI 使用不受缩放时间影响的动画。",
        "左侧显示根 茎 花三个槽位，右侧只显示与当前槽位兼容的已收集部件。",
        "选择部件时预览 新能力 替换掉的能力 外观缩略图。按确认后触发 GraftApplied 事件。",
        "角色控制器 战斗模块 HUD 外观挂点和音频监听同一事件刷新，避免各系统自行读取导致状态不一致。",
        "关闭界面后保留 0.25 秒输入保护，防止确认键同时触发普攻或技能。",
    ],
)

add_heading(doc, "5.4 嫁接验收", 2)
add_bullets(
    doc,
    [
        "敌人仍在场时可以打开嫁接界面，暂停期间玩家和敌人位置不变化。",
        "装上藤蔓触须后，站在默认攻击范围外仍可命中，且一次可命中前后两只敌人。",
        "装上毒菌伞后，技能不再生成种子投射物，而是生成有持续伤害和减速的毒雾区。",
        "装上铁甲根后，同一伤害源从 12 点降为 9 点，冲刺护盾可以消除蘑菇的普通毒球。",
        "卸下或替换部件后，旧效果立即移除，不残留减伤 中毒区生成逻辑或视觉挂点。",
    ],
)

add_page_break(doc)
add_heading(doc, "6 敌怪设计", 1)
add_body(doc, "三个敌人分别教会玩家处理贴身压力、远程火力和蓄力冲刺。所有 AI 使用小型有限状态机，不做复杂平台寻路；敌人在各自战斗区内活动，越界时返回。")

add_heading(doc, "6.1 数值总表", 2)
add_table(
    doc,
    ["敌人", "生命", "伤害", "移动", "攻击节奏", "掉落"],
    [
        ["腐化藤蔓怪", "25", "8", "1.8", "进入 1.2 单位后前摇 0.45 秒抽打；冷却 1.5 秒", "藤蔓触须首杀必掉"],
        ["毒孢蘑菇", "18", "6", "0.6", "6 单位内前摇 0.65 秒发射毒球；冷却 2.5 秒", "毒菌伞首杀必掉"],
        ["废铁甲虫", "35", "12", "3.2 冲刺", "瞄准 0.8 秒后直线冲刺；撞墙眩晕 1.0 秒；冷却 1.8 秒", "铁甲根必掉"],
    ],
    widths=[1.25, 0.6, 0.6, 0.9, 2.7, 1.05],
    font_size=8.5,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "6.2 行为状态", 2)
add_table(
    doc,
    ["敌人", "状态序列", "玩家可读信息", "防卡死规则"],
    [
        ["藤蔓怪", "待机 追踪 前摇 攻击 后摇 受击 死亡", "攻击前身体后仰，藤蔓高亮并伴随短声", "离玩家过远则回出生点；台阶前停下"],
        ["蘑菇", "待机 锁定 前摇 发射 冷却 受击 死亡", "菌盖收缩后张开，发射线方向可预测", "无视小平台寻路；保持原地或短距离左右移动"],
        ["甲虫", "巡逻 锁定 蓄力 冲刺 撞墙眩晕 受击 死亡", "地面出现窄长警示带，甲壳闪烁并发出摩擦声", "冲刺有最大距离；越界立即停下；撞墙必进入眩晕"],
    ],
    widths=[1.05, 2.5, 2.15, 1.4],
    font_size=8.6,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "6.3 敌怪通用规则", 2)
add_bullets(
    doc,
    [
        "同一时刻最多 4 只敌人激活攻击意图，避免狭窄横版画面出现无法读秒的围攻。",
        "敌人死亡在逻辑上立即从存活计数移除，尸体动画 0.6 至 1.0 秒后回收。",
        "投射物离开相机外扩边界或存在超过 3 秒后销毁，避免清场条件被无关物体影响。",
        "敌人受到毒雾减速时，攻击前摇不缩短；甲虫冲刺状态免疫减速，蓄力状态可被减速。",
        "最后一只敌人死亡后停止生成新攻击与投射物，再触发清场流程。",
    ],
)

add_page_break(doc)
add_heading(doc, "7 程序交付清单", 1)
add_body(
    doc,
    "程序工作以完成闭环为准，采用少量清晰模块和 ScriptableObject 配置，不在 48 小时内搭建通用 Roguelite 框架。"
    "融合 配方 随机房间和永久存档不写半成品界面，也不进入主流程。",
)

add_heading(doc, "7.1 场景与构建", 2)
add_table(
    doc,
    ["交付物", "内容", "完成标准"],
    [
        ["Boot", "初始化输入 音频 分辨率和一次运行数据", "不依赖编辑器状态，可直接进入 MainMenu"],
        ["MainMenu", "开始 操作说明 音量 退出", "键鼠可完整操作，按钮无重复触发"],
        ["Level01", "Tilemap 场景 三战斗区 嫁接 传送门", "从载入到通关无场景重载和软锁"],
        ["Result", "完成用时 已嫁接部件 重玩 返回主页面", "重玩会清空局内库存和敌人状态"],
        ["Windows Build", "x86 64 发布包 README 和 credits", "在无 Unity 环境的第二台电脑启动并完成一局"],
    ],
    widths=[1.2, 3.15, 2.75],
    font_size=9,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "7.2 核心模块", 2)
add_table(
    doc,
    ["模块", "职责", "建议组件或数据", "P0 验收"],
    [
        ["输入", "统一键鼠和可选手柄动作", "Unity Input System InputActionAsset", "场景切换后不重复订阅"],
        ["角色控制", "移动 跳跃 冲刺 朝向 地面检测", "PlayerController2D PlayerState", "斜坡边缘不抖动，死亡后禁用"],
        ["生命与伤害", "伤害 无敌 击退 死亡", "Health DamageInfo Hurtbox Hitbox", "同一攻击不重复命中"],
        ["攻击", "默认普攻 种子 毒雾 藤鞭 铁甲冲刺", "PlayerCombat AttackData Projectile Pool", "嫁接后行为即时切换"],
        ["状态效果", "持续伤害 减速 临时免伤", "StatusEffectController", "毒雾刷新规则一致，不无限叠层"],
        ["敌人 AI", "三种有限状态机与边界", "EnemyBrain EnemyData", "前摇可读，越界可恢复"],
        ["战斗区", "封锁 出怪 存活计数 清场", "EncounterManager EncounterZone", "最后敌人死亡只结算一次"],
        ["掉落与库存", "关键首掉 拾取 部件库存", "LootDrop RunInventory TraitData", "漏拾可补发，重开清空"],
        ["嫁接", "槽位切换 效果应用 外观刷新", "GraftManager GraftPanel GraftApplied", "旧效果完整移除"],
        ["传送门", "监听清场 开启动画 交互结算", "ExitPortal LevelFlow", "未清场不可交互"],
        ["UI", "HUD 提示 嫁接 暂停 死亡 完成", "UIRoot ViewModel 或事件绑定", "分辨率变化不遮挡"],
        ["音频", "音乐状态和事件音效", "AudioManager AudioMixer", "音量可调，场景切换不叠播"],
    ],
    widths=[0.85, 1.65, 2.55, 2.05],
    font_size=8.1,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "7.3 数据资产", 2)
add_table(
    doc,
    ["ScriptableObject", "关键字段", "数量"],
    [
        ["PlayerStatsData", "maxHp moveSpeed jumpVelocity dashDistance dashCooldown invincibleTime", "1"],
        ["AttackData", "id damage shape range startup active recovery cooldown knockback prefab", "默认普攻 种子 藤鞭 毒雾 共 4"],
        ["EnemyData", "id maxHp moveSpeed damage attackRange cooldown dropId", "3"],
        ["TraitData", "id slot displayName description statModifiers attackOverride visualPrefab icon audioEvent", "3"],
        ["EncounterData", "enemyPrefab spawnPoint delay guaranteedDrop", "3 个战斗区"],
    ],
    widths=[1.65, 4.25, 1.2],
    font_size=8.7,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
)
add_body(
    doc,
    "TraitData 可保留 tags 和 effectList 字段，供后续融合读取，但本次不创建 FusionRecipeData，不实现配方检查。所有 ID 在启动时检查重复和空引用，"
    "错误要包含资产名，便于现场修复。",
)

add_heading(doc, "7.4 程序目录建议", 2)
add_table(
    doc,
    ["目录", "内容"],
    [
        ["Assets Game Scripts Core", "GameFlow EventBus Pool 基础接口"],
        ["Assets Game Scripts Player", "控制 战斗 生命 嫁接应用"],
        ["Assets Game Scripts Enemies", "三种敌人状态与通用 EnemyBrain"],
        ["Assets Game Scripts Level", "EncounterZone 生成点 传送门"],
        ["Assets Game Scripts UI", "主页面 HUD 嫁接 暂停 结算"],
        ["Assets Game Data", "Player Enemy Attack Trait Encounter 配置"],
        ["Assets Game Art Audio Prefabs Scenes", "按资产类型分组，避免 Resources 无序加载"],
    ],
    widths=[2.55, 4.55],
    font_size=9,
    alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_page_break(doc)
add_heading(doc, "8 美术交付清单", 1)
add_body(
    doc,
    "美术目标是高可读的手绘 2D 或像素手绘混合风格。首版用灰绿和焦褐表现枯萎环境，用清亮绿黄识别玩家，用紫红或暖红标记危险。"
    "攻击预警不能只依赖颜色，必须同时有轮廓或形状变化。",
)

add_heading(doc, "8.1 主角与嫁接资产", 2)
add_table(
    doc,
    ["资产", "最低动画或状态", "数量与规格", "验收"],
    [
        ["主角基础体", "待机 跑 跳上升 下落 冲刺 普攻 技能 受击 死亡", "1 套；角色画面高度约 160 至 192 px；统一锚点在脚底", "所有动作方向和命中帧一致，不明显穿地"],
        ["藤蔓触须", "待机挂点 普攻伸展 收回", "茎部覆盖层 1 套；藤鞭轨迹 1 套", "攻击长度和程序判定一致"],
        ["毒菌伞", "待机挂点 技能释放轻摆", "花部覆盖层 1 套；毒雾区域 1 套", "毒雾边界和持续时间可读"],
        ["铁甲根", "待机挂点 跑动 冲刺护盾", "根部覆盖层 1 套；挡弹碎裂 1 套", "护盾只在有效窗口点亮"],
    ],
    widths=[1.2, 2.3, 2.25, 1.35],
    font_size=8.5,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "8.2 敌人与场景资产", 2)
add_table(
    doc,
    ["资产包", "交付内容", "建议数量", "P0 验收"],
    [
        ["藤蔓怪", "待机 移动 攻击 受击 死亡；藤蔓攻击特效", "1 角色 5 动画", "前摇与攻击帧对齐"],
        ["毒孢蘑菇", "待机 小移 发射 受击 死亡；毒球和命中特效", "1 角色 5 动画 2 特效", "发射方向清晰，弹体不融入背景"],
        ["废铁甲虫", "巡逻 蓄力 冲刺 眩晕 受击 死亡；地面警示带", "1 角色 6 动画 1 预警", "蓄力和冲刺状态一眼可分"],
        ["枯萎林地", "地面 平台 边缘 阴影 斜坡或台阶；3 层背景；前景装饰", "12 至 16 个 Tile；6 个可复用道具；3 层背景", "可站立面与纯装饰明确分层"],
        ["战斗封锁", "枯根墙关闭 开启两个状态", "1 套", "战斗中阻挡，清场后消失或缩回"],
        ["传送门", "枯死 破土生长 开启 待机 进入", "1 套；至少 5 个状态", "未开启和可进入状态差异明显"],
        ["拾取物", "三种器官图标 地面轮廓光 拾取粒子", "3 图标 3 小图 1 通用粒子", "小尺寸仍能识别槽位"],
    ],
    widths=[1.15, 2.8, 1.8, 1.35],
    font_size=8.35,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "8.3 战斗特效", 2)
add_table(
    doc,
    ["特效", "用途", "最低要求"],
    [
        ["普通挥击", "默认普攻轨迹", "持续不超过 0.2 秒，不遮敌人"],
        ["藤鞭轨迹", "嫁接后的长距离普攻", "长度与判定一致，穿透时显示第二段命中"],
        ["种子弹与命中", "默认技能", "弹体轮廓清楚，命中后立刻消失"],
        ["毒雾", "持续伤害和减速区域", "边界明确，中心半透明，结束前渐隐"],
        ["冲刺残影", "基础位移", "2 至 3 层短残影"],
        ["铁甲护盾与碎裂", "减伤窗口和挡弹", "有效期亮起，挡弹单独爆点"],
        ["受击和死亡", "玩家与敌人通用", "闪白或染色加小粒子，死亡不遮挡掉落"],
        ["清场和门开启", "关卡高潮", "根系生长尘土 花粉和门芯亮起"],
    ],
    widths=[1.45, 2.3, 3.35],
    font_size=9,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "8.4 UI 资产", 2)
add_table(
    doc,
    ["界面", "美术交付"],
    [
        ["主页面", "标题字 1；背景 1；按钮普通 悬停 按下 3 态；音量图标；键位页底板"],
        ["HUD", "生命条底与填充；技能冷却遮罩；根 茎 花槽框；3 部件图标；交互键提示"],
        ["嫁接", "三槽人体示意；部件列表项；选中态；效果对比箭头；确认 返回按钮"],
        ["结算", "死亡与完成标题；用时与部件汇总底板；重玩和主页面按钮"],
    ],
    widths=[1.25, 5.85],
    font_size=9,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "8.5 美术导入规范", 2)
add_bullets(
    doc,
    [
        "透明资产使用 PNG，统一 Pixels Per Unit；若用像素风，统一 Point Filter 和同一网格尺寸。",
        "命名采用 类别 对象 状态 序号，例如 chr spirit attack 01 和 vfx poison cloud 01。",
        "主角器官使用固定挂点 Root Stem Flower，所有覆盖层以同一脚底原点导出。",
        "Sorting Layer 至少拆分 Background Terrain Actor VFX Foreground UI，危险预警位于角色与地形之间。",
        "每个动画在交付表标记命中帧 发射帧 无敌开始和结束帧，程序按事件点绑定。",
    ],
)

add_page_break(doc)
add_heading(doc, "9 音乐与音效交付清单", 1)
add_body(
    doc,
    "音频服务于动作反馈和清场回报。音乐不必数量多，但循环必须自然，且不能盖住敌方前摇。首版使用 Unity Audio Mixer 分为 Music SFX UI 三组，"
    "至少提供主音量控制。",
)

add_heading(doc, "9.1 音乐", 2)
add_table(
    doc,
    ["曲目", "长度与结构", "情绪与用途", "交付格式"],
    [
        ["主页面音乐", "30 至 45 秒无缝循环；可有 2 秒引子", "脆弱幼芽和未知森林，节奏克制", "WAV 48 kHz 24 bit；另交循环点说明"],
        ["第一关音乐", "75 至 100 秒无缝循环", "枯萎森林的木质打击和轻节奏；不做战斗分层也可", "WAV 48 kHz 24 bit"],
        ["清场短句", "3 至 5 秒，不循环", "从压抑转为明亮，衔接传送门开启", "WAV 48 kHz 24 bit"],
        ["完成短句", "4 至 6 秒，不循环", "明确结束与希望感", "WAV 48 kHz 24 bit"],
    ],
    widths=[1.2, 2.1, 2.5, 1.3],
    font_size=8.8,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "9.2 音效事件", 2)
add_table(
    doc,
    ["类别", "必须交付的事件", "建议变体数"],
    [
        ["玩家动作", "jump land dash attack swing attack hit seed shoot seed hit hurt death", "命中和挥击各 2 至 3 个，其余 1 个"],
        ["嫁接能力", "graft pickup graft confirm vine swing poison cast poison loop iron guard projectile break", "确认 2 个；其他各 1 个"],
        ["敌人", "vine telegraph vine attack mushroom telegraph mushroom shoot beetle charge beetle impact enemy hurt enemy death", "受击和死亡各 2 个，其余 1 个"],
        ["流程", "encounter lock encounter clear portal grow portal idle portal enter", "各 1 个；portal idle 可循环"],
        ["UI", "hover click back error pause result", "hover click 各 2 个，其余 1 个"],
    ],
    widths=[1.05, 4.85, 1.2],
    font_size=8.6,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "9.3 混音与验收", 2)
add_bullets(
    doc,
    [
        "音乐循环点无明显爆音或静音缝，进入第一关时做 0.5 至 1.0 秒淡入淡出。",
        "敌方攻击前摇、玩家受击、嫁接确认和传送门开启在音乐存在时仍清楚可辨。",
        "所有源文件保留无损版本，Unity 内可按需要压缩；循环环境音使用 Vorbis，短音效优先 PCM 或 ADPCM。",
        "任何素材若来自外部库，必须确认允许 Game Jam 和公开发布，并在 credits.txt 写明名称 作者 来源和许可证。",
        "总线不削波。音乐基准约为负 16 至负 14 LUFS，音效按听感分层，最终以实际游戏场景校准。",
    ],
)

add_page_break(doc)
add_heading(doc, "10 团队接口与交付规范", 1)

add_heading(doc, "10.1 程序与美术接口", 2)
add_table(
    doc,
    ["接口项", "程序提供", "美术提供", "冻结时间"],
    [
        ["角色尺寸", "碰撞体和 PPU 基准预制体", "按脚底原点导出的主角与覆盖层", "第 4 小时"],
        ["动画事件", "attack hit skill fire dash iframe 事件名", "动画表标出对应帧", "第 8 小时"],
        ["挂点", "Root Stem Flower 三个 Transform", "每个器官相同坐标原点", "第 8 小时"],
        ["敌人判定", "Hitbox Hurtbox Gizmo 截图", "攻击轮廓和前摇动画", "第 12 小时"],
        ["传送门", "Closed Opening Open Enter 状态接口", "对应动画和粒子序列", "第 20 小时"],
        ["UI", "1920 x 1080 锚点线框", "九宫格底板 图标 按钮状态", "第 18 小时"],
    ],
    widths=[1.1, 2.35, 2.6, 1.05],
    font_size=8.7,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
)

add_heading(doc, "10.2 程序与音频接口", 2)
add_body(
    doc,
    "音频文件名与事件 ID 一致。程序在事件发生处只调用 AudioManager.Play(eventId)，不由动画或敌人脚本直接查找 AudioClip。"
    "循环声音必须显式开始和停止，毒雾销毁 传送门离场和场景切换时均做清理。",
)

add_heading(doc, "10.3 版本控制", 2)
add_bullets(
    doc,
    [
        "主分支始终保持可运行；功能在短分支完成后合并。每次合并前打开 MainMenu 和 Level01 做 2 分钟冒烟测试。",
        "Unity 场景和 ProjectSettings 指定一名程序负责人合并，避免多人同时修改同一 YAML 文件。",
        "大贴图和 WAV 使用 Git LFS。Library Temp Logs 和本地构建目录不提交。",
        "第 32 小时起冻结目录和命名；第 40 小时起只允许修复阻断性问题，不再加入新功能。",
    ],
)

add_page_break(doc)
add_heading(doc, "11 48 小时制作排期", 1)
add_body(
    doc,
    "排期按至少一名程序 一名美术 一名音频或综合成员并行设计。若人员更少，优先减少动画变体和第三种敌人的独立外观，"
    "但不能删掉随时嫁接 攻击变化和清场传送门。",
)

add_table(
    doc,
    ["时间", "程序", "美术", "音乐音效", "共同里程碑"],
    [
        ["0 至 2 小时", "建立工程 场景 Git 忽略和输入", "确定角色轮廓 配色 PPU", "确定音乐动机和事件表", "冻结 P0 与本表数值"],
        ["2 至 8 小时", "移动 跳跃 冲刺 生命 普攻灰盒", "主角基础体和关键动作草稿", "主页面循环草稿；动作基础音效", "第 8 小时可在灰盒中移动和打木桩"],
        ["8 至 16 小时", "三敌人共用 FSM；第一关灰盒；战斗区", "3 敌人轮廓 攻击前摇；Tile 第一版", "第一关音乐草稿；敌人前摇音效", "第 16 小时可清三战斗区"],
        ["16 至 24 小时", "掉落 库存 嫁接界面 三能力切换", "3 嫁接覆盖层；藤鞭 毒雾 护盾", "嫁接 能力和拾取音效", "第 24 小时完成核心嫁接闭环"],
        ["24 至 32 小时", "传送门 结算 重开 HUD 音频接入", "传送门 UI 主页面 场景完善", "清场 传送门 完成短句；混音初版", "第 32 小时全流程可从头到尾"],
        ["32 至 38 小时", "Bug 修复 输入保护 对象池 性能", "受击 命中 粒子 背景层和清晰度", "补缺失事件 调音乐循环", "第 38 小时内容冻结"],
        ["38 至 44 小时", "多分辨率测试 构建 死亡和重开压力测", "修穿模 遮挡 预警问题", "实机混音和响度修正", "至少 3 名外部玩家完整试玩"],
        ["44 至 48 小时", "只修 P0；制作最终 Windows 包", "封面 截图 最终导入检查", "导出无损源文件和 credits", "第 46 小时提交候选，第 48 小时留作缓冲"],
    ],
    widths=[0.8, 1.75, 1.6, 1.55, 1.4],
    font_size=7.65,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "11.1 砍需求顺序", 2)
add_numbered(
    doc,
    [
        "先砍手柄 深度设置 伤害数字 镜头震动和复杂 UI 动画。",
        "再砍第三个战斗区的额外敌人数量，保留甲虫和铁甲根展示。",
        "再把敌人动画降为关键姿态加程序插值，保留前摇 攻击 受击 死亡四个可读状态。",
        "若主角动画来不及，保留待机 跑 跳 普攻 技能 受击六组，冲刺使用残影和拉伸处理。",
        "不得砍除随时嫁接 藤鞭攻击替换 毒雾技能替换 清场传送门和结算。它们共同构成作品身份。",
    ],
)

add_heading(doc, "11.2 时间闸门", 2)
add_table(
    doc,
    ["闸门", "必须达到", "未达到时立即处理"],
    [
        ["第 8 小时", "角色可稳定移动 跳跃 攻击和受伤", "停止菜单和美化，所有人支援控制与判定"],
        ["第 16 小时", "三战斗区灰盒可清完", "用同一敌人骨架做三种行为，取消独立复杂动画"],
        ["第 24 小时", "至少藤蔓和毒菌两种嫁接可即时切换", "停止所有 P1，先完成事件链和旧效果移除"],
        ["第 32 小时", "主页面至完成页可完整跑通", "冻结新内容，传送门和结算优先于所有表现"],
        ["第 40 小时", "候选构建可在第二台电脑运行", "只做崩溃 软锁 输入和资源丢失修复"],
    ],
    widths=[1.0, 3.15, 2.95],
    font_size=8.9,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_page_break(doc)
add_heading(doc, "12 验收与测试", 1)

add_heading(doc, "12.1 完整流程验收", 2)
add_table(
    doc,
    ["编号", "测试步骤", "通过标准"],
    [
        ["F01", "启动 Windows 构建并点击开始", "5 秒内进入第一关，无空白或重复加载"],
        ["F02", "完成移动 跳跃 冲刺 普攻 技能", "输入可响应，角色不穿地，不永久锁定"],
        ["F03", "击杀战斗区一首只藤蔓怪", "必定出现藤蔓触须，拾取后进入库存"],
        ["F04", "敌人存活时打开嫁接界面", "游戏暂停，敌人 投射物和计时停止"],
        ["F05", "嫁接藤蔓触须后攻击两名排成直线的敌人", "攻击距离变长并穿透一名敌人，外观和音效同步变化"],
        ["F06", "嫁接毒菌伞并释放技能", "不生成种子弹；生成持续 3 秒的毒雾并造成持续伤害"],
        ["F07", "嫁接铁甲根后承受 12 点攻击", "只损失 9 点生命；冲刺可消除普通毒球"],
        ["F08", "来回替换同槽部件", "旧效果和旧挂点被移除，不重复叠加"],
        ["F09", "留下任意敌人靠近出口", "传送门不可用，不出现完成提示"],
        ["F10", "击杀最后一只敌人", "0.8 至 1.5 秒内播放清场反馈并开启传送门，仅触发一次"],
        ["F11", "进入传送门", "锁定输入并进入完成页，用时和部件记录正确"],
        ["F12", "从完成页重玩", "所有敌人 传送门 库存 生命和计时恢复初始状态"],
        ["F13", "玩家死亡并选择重试", "无重复死亡页，重试后流程可再次完成"],
        ["F14", "切换 1920x1080 1600x900 1280x720", "HUD 嫁接和菜单无裁切重叠"],
        ["F15", "连续游玩三局", "无明显内存增长 音乐叠播 丢失输入或存活计数错误"],
    ],
    widths=[0.6, 3.25, 3.25],
    font_size=8.35,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "12.2 角色分工完成定义", 2)
add_table(
    doc,
    ["岗位", "完成定义"],
    [
        ["程序", "第二台电脑上的最终构建可从主页面完成一局；F01 至 F15 中无阻断项；Console 无持续报错；重新开始三次状态正确"],
        ["美术", "清单内 P0 资产已导入；角色和器官不穿模；攻击和预警与判定对齐；三种敌人 三种部件 传送门均可在游戏画面中立即区分"],
        ["音乐音效", "两首循环音乐和关键事件音效全部接入；无爆音 叠播或明显音量跳变；敌方前摇和传送门开启在音乐中清楚可闻；授权记录完整"],
        ["策划或制作", "全流程 3 次外部试玩记录完成；首局平均不超过 8 分钟；至少 2 名玩家在无讲解情况下成功嫁接并通关；问题按 P0 P1 排序"],
    ],
    widths=[1.2, 5.9],
    font_size=9,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "12.3 提交包", 2)
add_bullets(
    doc,
    [
        "Windows 可执行文件与 Data 文件夹，目录名不含 临时 最终版2 等模糊词。",
        "README.txt，写明启动方式 操作键位 游戏目标 已知问题和团队成员。",
        "credits.txt，列出音乐 字体 插件 贴图和所有外部素材的来源与许可证。",
        "至少 3 张截图：主页面；嫁接界面；藤鞭或毒雾战斗；可额外加入传送门开启画面。",
        "30 至 60 秒演示视频，顺序为击杀掉落 随时嫁接 攻击变化 清场传送门。",
        "Unity 工程压缩包或版本库最终标签，确保项目可恢复。",
    ],
)

add_page_break(doc)
add_heading(doc, "13 风险与对策", 1)
add_table(
    doc,
    ["风险", "早期信号", "处理方式"],
    [
        ["嫁接只改数值", "试玩者看不出是否装备成功", "每个部件至少绑定一个行为替换和一个外观或特效变化；以录屏验收"],
        ["随时嫁接导致战斗误操作", "打开界面后仍被攻击或确认键触发技能", "嫁接界面暂停时间；关闭后加 0.25 秒输入保护"],
        ["掉落随机导致核心未展示", "测试局未获得毒菌伞或铁甲根", "三种关键部件首杀必掉，出口旁补发漏拾物"],
        ["敌人死亡计数软锁", "场上无敌人但门不开", "敌人生成时注册，死亡逻辑先注销；提供开发者强制清场键仅用于调试"],
        ["器官图层穿模", "跑跳时根 茎 花错位", "冻结三个挂点和统一原点；优先用覆盖层，不做自由骨骼重组"],
        ["美术量超时", "第 16 小时敌人仍无攻击前摇", "减少帧数和背景道具，保留关键姿态与特效，不增加第四种敌人"],
        ["音频最后接入", "第 32 小时仍用临时文件名", "第 4 小时冻结事件名，音频按 ID 交付，程序可批量替换"],
        ["范围继续膨胀", "有人开始做融合 Boss 随机门或剧情", "制作人按 P0 表拒绝合并，新增想法只进入后续清单"],
    ],
    widths=[1.55, 2.5, 3.05],
    font_size=8.6,
    alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "13.1 数值调优顺序", 2)
add_numbered(
    doc,
    [
        "先调角色移动 跳跃 攻击前后摇和命中反馈，再调敌人血量。",
        "确保默认攻击击杀藤蔓怪约需 3 次，藤鞭以范围和穿透取胜而不是纯伤害碾压。",
        "确保毒雾完整覆盖可击杀蘑菇但不能瞬间清空混编，玩家需要把敌人引入区域。",
        "确保甲虫冲刺可以躲避且撞墙后有稳定输出窗口，再调伤害。",
        "最后调整关卡敌人数和治疗量，使首次游玩 5 至 8 分钟、单次死亡不超过一次为合理目标。",
    ],
)

add_page_break(doc)
add_heading(doc, "14 后续优化方向", 1)
add_body(
    doc,
    "后续开发沿现有闭环向外扩展，先深化嫁接，再加入融合台，最后扩展 Roguelite 路线。任何新系统都必须继续满足 身体外观改变和操作结果改变 两个条件。",
)

add_heading(doc, "14.1 第一阶段 融合台与第二关", 2)
add_table(
    doc,
    ["内容", "规则", "复用首版成果"],
    [
        ["融合台", "只出现在安全区或关卡间；选择两个已嫁接特征；预览产物和代价；确认后生成复合特征", "复用嫁接界面 槽位 TraitData 和 GraftApplied 刷新链"],
        ["瘟疫菌丝", "藤蔓触须加毒菌伞；藤鞭命中附毒，毒雾击杀向附近敌人扩散一次；技能冷却增加 1 秒", "复用藤鞭 毒雾 中毒和穿透"],
        ["荆棘古臂", "铁甲根加后续树人之臂；受击后下一次重击释放刺波；移动速度略降", "复用减伤 护盾 击退和范围命中"],
        ["第二关", "新增一种环境机制 两种敌人 一个小型精英；通关继续使用传送门", "复用 EncounterManager Portal 和全部角色系统"],
    ],
    widths=[1.35, 3.75, 2.0],
    font_size=8.7,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)
add_body(
    doc,
    "融合必须在融合台完成，不能像嫁接一样随时操作。这样可以让嫁接承担高频调整，融合承担关卡间的重大构筑承诺，也能给专属形态 动画 配方说明和风险代价留出清晰展示空间。",
)

add_heading(doc, "14.2 第二阶段 首章垂直切片", 2)
add_table(
    doc,
    ["方向", "目标"],
    [
        ["根门推进", "清场后从两扇门选择奖励方向；同一随机种子可复现；所有路线可到 Boss"],
        ["内容扩展", "3 至 5 个手工房间模块；5 种小怪；1 个精英；根 茎 花共 6 至 8 个特征"],
        ["成长选择", "每 2 个房间一次三选一进化，优先出现与当前器官标签相关的选项"],
        ["Boss", "腐化古树至少 3 种可学习招式；根刺 腐果 近身旋转；战斗时长 90 至 150 秒"],
        ["净化反馈", "Boss 击败后同一场景从灰绿枯败切换为湿润苔藓 暖光和新芽，音乐同步转调"],
    ],
    widths=[1.45, 5.65],
    font_size=9,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "14.3 第三阶段 完整 Roguelite", 2)
add_bullets(
    doc,
    [
        "开放叶 果实 种子槽位，加入滑翔 召唤 扩散 储能等移动与战斗变化。",
        "加入污染值和高风险房间，但自然路线始终可以通关，污染不成为唯一最优解。",
        "扩展随机房间权重 保底和可达性标签，加入 seed 回放用于 Bug 复现。",
        "加入图鉴 基地 NPC 配方线索和以解锁选择池为主的永久成长，避免纯数值碾压。",
        "完善键鼠与手柄 可访问性 画面设置 本地存档和版本迁移。",
        "继续扩展融合，但每个融合都必须有独立循环 外观和代价，不做仅加百分比的伪融合。",
    ],
)

add_heading(doc, "14.4 后续开发顺序", 2)
add_table(
    doc,
    ["顺序", "里程碑", "进入条件"],
    [
        ["1", "修复 Jam 反馈并稳定首版", "10 名试玩者中至少 8 名能无讲解完成嫁接和通关"],
        ["2", "加入融合台和 2 条配方", "嫁接状态切换无残留，特征数据可组合"],
        ["3", "加入第二关和 Boss 原型", "传送门与 Encounter 可跨场景复用"],
        ["4", "加入根门路线和 3 至 5 个房间模块", "固定流程已达到稳定手感和数值基线"],
        ["5", "加入进化 污染 图鉴和永久解锁", "核心战斗与构筑已经证明有重玩价值"],
    ],
    widths=[0.7, 2.5, 3.9],
    font_size=9,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_page_break(doc)
heading = add_heading(doc, "15 最终开工清单", 1)
heading.paragraph_format.page_break_before = True
add_body(doc, "团队开工前只需要确认以下项目。确认后 48 小时内不再讨论新增系统，只围绕可玩闭环和反馈质量决策。")
add_table(
    doc,
    ["确认项", "最终值"],
    [
        ["一句话玩法", "击败敌怪取得器官，随时嫁接并立即改变攻击，清场后进入传送门"],
        ["场景", "MainMenu Level01 Result 三个可见页面或场景，Boot 可后台存在"],
        ["敌人", "藤蔓怪 蘑菇 甲虫共 3 种"],
        ["嫁接", "藤蔓触须改普攻 毒菌伞改技能 铁甲根改防御与冲刺"],
        ["掉落", "关键首杀必掉，无随机掉率"],
        ["第一关", "教学起点加 3 战斗区加传送门出口"],
        ["通关", "全敌人清除后门开启，进入后显示完成页"],
        ["不做", "融合 Boss 随机关卡 污染 进化 商店 图鉴 永久成长；第 38 小时内容冻结"],
    ],
    widths=[1.3, 5.8],
    font_size=9.2,
    alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
